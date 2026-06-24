"""
Tool implementations for the Tender Optimization assistant.

These are pure functions over pandas DataFrames — no Streamlit, no Bedrock — so
they can be unit-tested and adversarially probed in isolation. The chat layer
(`chat_ui.py`) wires them to the Bedrock Converse tool-use loop.

The assistant works against two artifacts:
  * comprehensive_data: the merged GVT+rate+performance table the dashboard builds.
    Key columns: Week Number, Category, SSL, Vessel, Discharged Port,
    Dray SCAC(FL), Facility, Terminal, Lane, Container Count, Base Rate,
    Total Rate, CPC, Performance_Score.
  * constraints: a list of dicts, one per row, using the constraint-file schema
    (see CONSTRAINT_COLUMNS) that `constraints_processor.process_constraints_file`
    expects.
"""
from __future__ import annotations

import io
import re
from typing import Optional

import pandas as pd

from ..core.utils import parse_day_of_week

# The exact constraint-file schema understood by constraints_processor.py.
CONSTRAINT_COLUMNS = [
    "Category", "Carrier", "Lane", "Port", "Week Number", "Day of Week",
    "Terminal", "SSL", "Vessel", "Maximum Container Count",
    "Minimum Container Count", "Percent Allocation", "Excluded FC", "Priority Score",
]

# Columns the user may filter / scope a constraint by (everything except amounts).
_SCOPE_FIELDS = ["Category", "Carrier", "Lane", "Port", "Week Number", "Day of Week",
                 "Terminal", "SSL", "Vessel"]
_AMOUNT_FIELDS = ["Maximum Container Count", "Minimum Container Count", "Percent Allocation"]

# Map a constraint scope field -> the column it filters in comprehensive_data.
_DATA_COLUMN_FOR = {
    "Category": "Category",
    "Carrier": "Dray SCAC(FL)",
    "Lane": "Lane",
    "Port": "Discharged Port",
    "Week Number": "Week Number",
    "Day of Week": "Day of Week",
    "Terminal": "Terminal",
    "SSL": "SSL",
    "Vessel": "Vessel",
}


# ==================== helpers ====================

def _carrier_col(df: pd.DataFrame) -> str:
    return "Dray SCAC(FL)" if "Dray SCAC(FL)" in df.columns else "Carrier"


def _coerce_num(val):
    try:
        if val is None or val == "":
            return None
        return float(val)
    except (TypeError, ValueError):
        return None


def _is_blank(val) -> bool:
    if val is None:
        return True
    if isinstance(val, float) and pd.isna(val):
        return True
    if isinstance(val, str) and val.strip() == "":
        return True
    return False


def _plain(v):
    """Coerce a possibly numpy/pandas scalar into a plain JSON-safe value.

    The constraint summary is built from DataFrame cells, so its numbers can be
    numpy ints/floats that the Bedrock tool-result serializer chokes on. Whole
    floats are normalized to int for cleaner output (priority 90.0 -> 90).
    """
    if v is None or isinstance(v, (bool, str)):
        return v
    if isinstance(v, (int, float)):
        if isinstance(v, float) and v.is_integer():
            return int(v)
        return v
    item = getattr(v, "item", None)  # numpy / pandas scalars expose .item()
    if callable(item):
        try:
            return _plain(item())
        except Exception:
            pass
    return str(v)


def _plain_scope_val(v):
    """Coerce a scope value (scalar or list, e.g. Excluded Facilities)."""
    if isinstance(v, list):
        return [_plain(x) for x in v]
    return _plain(v)


# ==================== analysis ====================

def analyze_data(df: Optional[pd.DataFrame], query_type: str = "overview",
                 group_by: Optional[str] = None, top_n: int = 10) -> dict:
    """Summarize the current dashboard data.

    query_type:
      - "overview": totals, carrier/lane/week counts, cost summary.
      - "by_carrier" / "by_lane" / "by_port" / "by_category" / "by_week":
        per-group container counts, total cost, avg rate, avg performance.
      - "cheapest" / "most_expensive": carriers ranked by average base rate.
      - "performance": carriers ranked by volume-weighted performance score.
    Returns a JSON-serializable dict.
    """
    if df is None or len(df) == 0:
        return {"error": "No data is loaded. Upload GVT and Rate files first."}

    carrier_col = _carrier_col(df)
    has_perf = "Performance_Score" in df.columns
    has_rate = "Base Rate" in df.columns
    total_containers = int(df["Container Count"].sum()) if "Container Count" in df.columns else len(df)

    if query_type == "overview":
        out = {
            "total_containers": total_containers,
            "total_rows": int(len(df)),
            "unique_carriers": int(df[carrier_col].nunique()) if carrier_col in df.columns else 0,
            "unique_lanes": int(df["Lane"].nunique()) if "Lane" in df.columns else 0,
        }
        if "Week Number" in df.columns:
            weeks = sorted(int(w) for w in df["Week Number"].dropna().unique())
            out["weeks"] = weeks
        if "Discharged Port" in df.columns:
            out["ports"] = sorted(str(p) for p in df["Discharged Port"].dropna().unique())
        if "Category" in df.columns:
            out["categories"] = sorted(str(c) for c in df["Category"].dropna().unique())
        if "Total Rate" in df.columns:
            out["total_cost"] = round(float(df["Total Rate"].sum()), 2)
        return out

    group_map = {
        "by_carrier": carrier_col, "by_lane": "Lane", "by_port": "Discharged Port",
        "by_category": "Category", "by_week": "Week Number",
    }
    if query_type in group_map:
        col = group_map[query_type]
        if col not in df.columns:
            return {"error": f"Column for '{query_type}' not available in data."}
        agg = {"Container Count": "sum"}
        if "Total Rate" in df.columns:
            agg["Total Rate"] = "sum"
        if has_rate:
            agg["Base Rate"] = "mean"
        if has_perf:
            agg["Performance_Score"] = "mean"
        grouped = df.groupby(col, dropna=False).agg(agg).reset_index()
        grouped = grouped.sort_values("Container Count", ascending=False).head(top_n)
        rows = []
        for _, r in grouped.iterrows():
            row = {"group": str(r[col]), "containers": int(r["Container Count"])}
            if "Total Rate" in grouped.columns:
                row["total_cost"] = round(float(r["Total Rate"]), 2)
            if "Base Rate" in grouped.columns:
                row["avg_base_rate"] = round(float(r["Base Rate"]), 2)
            if "Performance_Score" in grouped.columns and pd.notna(r["Performance_Score"]):
                row["avg_performance"] = round(float(r["Performance_Score"]), 3)
            rows.append(row)
        return {"query_type": query_type, "groups": rows}

    if query_type in ("cheapest", "most_expensive"):
        if not has_rate or carrier_col not in df.columns:
            return {"error": "Rate or carrier data not available."}
        valid = df[df["Base Rate"] > 0]
        if len(valid) == 0:
            return {"error": "No rows with a positive base rate."}
        agg = valid.groupby(carrier_col).agg(
            avg_base_rate=("Base Rate", "mean"),
            containers=("Container Count", "sum"),
        ).reset_index()
        agg = agg.sort_values("avg_base_rate", ascending=(query_type == "cheapest")).head(top_n)
        return {
            "query_type": query_type,
            "carriers": [
                {"carrier": str(r[carrier_col]),
                 "avg_base_rate": round(float(r["avg_base_rate"]), 2),
                 "containers": int(r["containers"])}
                for _, r in agg.iterrows()
            ],
        }

    if query_type == "performance":
        if not has_perf or carrier_col not in df.columns:
            return {"error": "Performance data not available."}
        agg = df.groupby(carrier_col).agg(
            avg_performance=("Performance_Score", "mean"),
            containers=("Container Count", "sum"),
        ).reset_index().sort_values("avg_performance", ascending=False).head(top_n)
        return {
            "query_type": "performance",
            "carriers": [
                {"carrier": str(r[carrier_col]),
                 "avg_performance": round(float(r["avg_performance"]), 3),
                 "containers": int(r["containers"])}
                for _, r in agg.iterrows()
            ],
        }

    return {"error": f"Unknown query_type '{query_type}'."}


# ==================== flip cost simulation ====================
#
# These thin wrappers build a read-only FlipSimulator over the current data and
# delegate to it. They exist so chat_ui's executor has a flat name->callable map
# and so the flip tools are unit-testable without Streamlit, like the rest of
# this module.

def _build_simulator(df: Optional[pd.DataFrame], rate_data=None, rate_type: str = "Base Rate"):
    from .simulation import FlipSimulator
    return FlipSimulator(df if df is not None else pd.DataFrame(), rate_data, rate_type)


def describe_selection(df, scope_input, rate_data=None, rate_type="Base Rate") -> dict:
    """Tool handler: summarize what a scope currently selects."""
    from .simulation import Scope
    sim = _build_simulator(df, rate_data, rate_type)
    return sim.describe_scope(Scope.from_dict(scope_input))


def simulate_flip(df, scope_input, target_carrier, rate_data=None, rate_type="Base Rate") -> dict:
    """Tool handler: price the scoped containers as if flipped to target_carrier."""
    from .simulation import Scope
    if _is_blank(target_carrier):
        return {"error": "No target carrier given. Ask the user which carrier to flip to."}
    sim = _build_simulator(df, rate_data, rate_type)
    return sim.simulate_flip(Scope.from_dict(scope_input), str(target_carrier))


def compare_carriers(df, scope_input, candidates, rate_data=None, rate_type="Base Rate") -> dict:
    """Tool handler: rank candidate carriers cheapest-first for the scoped containers."""
    from .simulation import Scope
    if isinstance(candidates, str):
        candidates = [candidates]
    candidates = [str(c) for c in (candidates or []) if not _is_blank(c)]
    if not candidates:
        return {"error": "No candidate carriers given to compare."}
    sim = _build_simulator(df, rate_data, rate_type)
    return sim.compare_carriers(Scope.from_dict(scope_input), candidates)


def lane_rate_options(df, scope_input, rate_data=None, rate_type="Base Rate") -> dict:
    """Tool handler: list rated carriers per lane for the scoped containers."""
    from .simulation import Scope
    sim = _build_simulator(df, rate_data, rate_type)
    return sim.lane_rate_options(Scope.from_dict(scope_input))


def flip_report(df, scope_input, target_carrier, rate_data=None,
                rate_type="Base Rate", max_rows=200) -> dict:
    """Tool handler: per-container old-vs-new carrier rate + savings report.

    The auditable, container-level companion to ``simulate_flip``. Mirrors the
    standalone Carrier Flip report's "GVT with New SCAC" sheet (Old Rate / New
    Rate / Savings per container) but reuses the same rate index the rest of the
    assistant prices flips with, so the numbers always agree.
    """
    from .simulation import Scope
    if _is_blank(target_carrier):
        return {"error": "No target carrier given. Ask the user which carrier to flip to."}
    try:
        cap = int(max_rows)
    except (TypeError, ValueError):
        cap = 200
    cap = max(1, min(cap, 2000))
    sim = _build_simulator(df, rate_data, rate_type)
    return sim.flip_report(Scope.from_dict(scope_input), str(target_carrier), max_rows=cap)


# ==================== optimization-aware tools ====================
#
# These reflect how the dashboard ACTUALLY decides allocations (a weighted
# cost+performance blend solved per [Lane, Week] group, with a historical-growth
# cap for supplier diversity) — not the naive re-pricing the flip tools do. The
# heavy lifting lives in the pure ``optimizer_core`` module; the configured
# weights are threaded in from the chat executor (which reads session state), so
# these handlers stay Streamlit-free and testable.


def get_optimization_settings(cost_weight: float, performance_weight: float,
                              max_growth_pct: float, n_historical_weeks: int = 5) -> dict:
    """Report the optimizer settings currently configured in the dashboard.

    The chat executor passes the live values from session state (0-1 floats), so
    the assistant can cite the REAL configured weights instead of assuming 70/30.
    """
    return {
        "cost_weight": round(float(cost_weight), 2),
        "performance_weight": round(float(performance_weight), 2),
        "cost_weight_pct": round(float(cost_weight) * 100),
        "performance_weight_pct": round(float(performance_weight) * 100),
        "max_growth_pct": round(float(max_growth_pct) * 100),
        "n_historical_weeks": int(n_historical_weeks),
        "objective": ("Per [Lane, Week] group the optimizer minimizes "
                      "cost_weight*normalized_cost + performance_weight*(1 - normalized_performance). "
                      f"Carriers are then capped at their last-{int(n_historical_weeks)}-week volume share "
                      f"+{round(float(max_growth_pct) * 100)}% growth, to preserve supplier diversity. "
                      "Carriers with no published rate on a lane get a 10x penalty rate (never treated as free)."),
    }


def recommend_carrier(df, scope_input, cost_weight: float, performance_weight: float,
                      max_growth_pct: float, top_n: int = 5) -> dict:
    """Tool handler: rank carriers for the scope using the optimizer's cost+perf blend."""
    from .optimizer_core import recommend_carriers_core
    try:
        n = int(top_n)
    except (TypeError, ValueError):
        n = 5
    n = max(1, min(n, 25))
    return recommend_carriers_core(
        df, scope_input,
        cost_weight=cost_weight, performance_weight=performance_weight,
        max_growth_pct=max_growth_pct, top_n=n,
    )


def run_optimization(df, scenario, scope_input=None, cost_weight: float = 0.7,
                     performance_weight: float = 0.3, max_growth_pct: float = 0.30,
                     historical_data=None, top_n: int = 25) -> dict:
    """Tool handler: run a dashboard reallocation scenario and compare to current.

    ``scenario`` is one of 'cheapest', 'performance', 'optimized'. Returns the
    current vs proposed cost, savings and %, per-carrier volume deltas, and the
    net containers reallocated. Read-only — the data is copied, never mutated.

    Weights are accepted as 0-1 floats (or 0-100, normalized here) so the chat
    executor can pass the live dashboard settings straight through.
    """
    from .optimizer_core import run_scenario_core
    try:
        cw, pw, mg = float(cost_weight), float(performance_weight), float(max_growth_pct)
    except (TypeError, ValueError):
        cw, pw, mg = 0.7, 0.3, 0.30
    if cw > 1 or pw > 1:  # tolerate 0-100 input
        total = (cw + pw) or 1.0
        cw, pw = cw / total, pw / total
    if mg > 1:
        mg = mg / 100.0
    try:
        n = max(1, min(int(top_n), 50))
    except (TypeError, ValueError):
        n = 25
    return run_scenario_core(
        df, scenario, scope_input,
        cost_weight=cw, performance_weight=pw, max_growth_pct=mg,
        historical_data=historical_data, top_n=n,
    )


def preview_optimization(df, constraints, rate_data, cost_weight: float,
                         performance_weight: float, max_growth_pct: float,
                         historical_data=None) -> dict:
    """Tool handler: run a constraint set through the optimizer and report the impact.

    ``constraints`` is the working-set list of constraint dicts (staged or applied);
    it is converted to the template DataFrame and run through the real pipeline on a
    copy. Returns current-vs-proposed cost / performance / carrier-mix deltas.
    """
    from .optimizer_core import constraint_impact_core
    rows = [c for c in (constraints or []) if not c.get("_problems")]
    if not rows:
        return {"error": "There are no valid constraints to preview. Draft and validate at "
                         "least one rule (generate_constraints) first."}
    constraints_df = constraints_to_dataframe(rows)
    return constraint_impact_core(
        df, constraints_df, rate_data,
        cost_weight=cost_weight, performance_weight=performance_weight,
        max_growth_pct=max_growth_pct, historical_data=historical_data,
    )


def optimization_summary(df, cost_weight: float, performance_weight: float,
                         max_growth_pct: float) -> dict:
    """Tool handler: configured weights + headline current-vs-optimized cost."""
    from .optimizer_core import cost_perf_rollup, run_allocation_core, MAX_OPT_ROWS
    settings = get_optimization_settings(cost_weight, performance_weight, max_growth_pct)
    if df is None or len(df) == 0:
        return {"settings": settings,
                "error": "No data is loaded, so there is no allocation to summarize."}
    if len(df) > MAX_OPT_ROWS:
        return {"settings": settings,
                "error": f"The loaded data has {len(df):,} rows — too many to optimize in chat. "
                         "Narrow the dashboard filters and try again."}
    current = cost_perf_rollup(df)
    try:
        optimized = run_allocation_core(
            df, cost_weight=cost_weight, performance_weight=performance_weight,
            max_growth_pct=max_growth_pct,
        )
    except Exception as e:  # noqa: BLE001
        return {"settings": settings, "current": current,
                "error": f"Optimization failed: {e}"}
    proposed = cost_perf_rollup(optimized)
    cur = current["total_cost"] or 0.0
    new = proposed["total_cost"] or 0.0
    return {
        "settings": settings,
        "current": current,
        "optimized": proposed,
        "cost_delta": round(new - cur, 2),
        "cost_delta_pct": round(100.0 * (new - cur) / cur, 1) if cur else None,
        "note": ("'Optimized' reapplies the cost+performance blend and growth cap to ALL loaded "
                 "containers. Cost-only comparisons can look cheaper but ignore performance and "
                 "supplier-diversity guardrails."),
    }


def apply_constraints(staged_constraints: Optional[list]) -> dict:
    """Validation half of the direct-apply tool — split the working set into apply/reject.

    Pure: this NEVER touches session state. It mirrors the Apply button's filter
    (only rows with no validation problems are applicable). The chat executor
    performs the actual ``st.session_state.chatbot_applied_constraints`` write.
    """
    staged = staged_constraints or []
    if not staged:
        return {"applied_count": 0, "to_apply": [], "rejected": [],
                "note": ("There are no staged constraints to apply. Draft a rule with "
                         "generate_constraints (or upload a constraint file) first.")}
    to_apply, rejected = [], []
    for i, c in enumerate(staged):
        if c.get("_problems"):
            rejected.append({"index": i, "problems": c["_problems"],
                             **{col: _plain(c.get(col)) for col in CONSTRAINT_COLUMNS
                                if not _is_blank(c.get(col))}})
        else:
            to_apply.append(c)
    return {
        "applied_count": len(to_apply),
        "to_apply": to_apply,
        "applied": [
            {col: _plain(c.get(col)) for col in CONSTRAINT_COLUMNS if not _is_blank(c.get(col))}
            for c in to_apply
        ],
        "rejected": rejected,
    }


# ==================== open-ended analysis (sandboxed) ====================

def run_analysis(df: Optional[pd.DataFrame], code, max_rows=200,
                 save_as: Optional[str] = None,
                 memory: Optional[dict] = None) -> dict:
    """Tool handler: run model-written pandas against a read-only copy of the data.

    This is the escape hatch for analytical questions the fixed tools can't
    express (custom pivots, multi-column groupings, derived metrics). The snippet
    must assign its answer to a ``result`` variable; ``df`` is the data. Execution
    is sandboxed (see ``code_sandbox`` for the security model) — model code cannot
    import, open files, touch the network, or read process credentials, and runs
    against a copy so it can never mutate the dashboard's data.

    ANALYSIS MEMORY (multi-turn): ``memory`` is a dict of named results from
    earlier turns, exposed to the snippet as a read-only ``memory`` dict so a
    later analysis can build on prior work (e.g. ``memory['lax_breakdown']``).
    When ``save_as`` is set, the success payload carries the RAW result under
    ``_save`` ({"name", "value"}) so the chat layer can store it in session memory
    for a future turn to recall — this function stays pure (no session state).

    IMPORTANT: numbers produced here come from arbitrary code, not the tested
    cost model — they are NOT subject to the unrated-lane / ambiguity guards that
    protect ``simulate_flip`` / ``flip_report``. For pricing a flip, prefer those.
    """
    from .code_sandbox import run_sandboxed_code, DEFAULT_MAX_ROWS
    if _is_blank(code) or not isinstance(code, str):
        return {"ok": False, "error": "No analysis code provided."}
    try:
        cap = int(max_rows)
    except (TypeError, ValueError):
        cap = DEFAULT_MAX_ROWS
    cap = max(1, min(cap, 2000))
    want_raw = not _is_blank(save_as)
    out = run_sandboxed_code(df, code, max_rows=cap, memory=memory, return_raw=want_raw)
    # Detach the raw object into a _save directive for the chat layer; never leak
    # the un-coerced object back to the model (it isn't JSON-serializable).
    raw = out.pop("raw", None)
    if out.get("ok") and want_raw and raw is not None:
        out["_save"] = {"name": str(save_as).strip(), "value": raw}
        out["saved_as"] = str(save_as).strip()
    return out


def list_analysis_memory(memory: Optional[dict] = None) -> dict:
    """List the named results saved in analysis memory this session.

    The multi-turn recall companion to ``run_analysis(save_as=...)``: returns a
    compact, JSON-safe catalogue (name, kind, shape/columns or value) of every
    stored result so the assistant can decide what to ``recall`` (read via the
    ``memory`` dict) in a follow-up snippet instead of recomputing.
    """
    memory = memory or {}
    if not memory:
        return {
            "count": 0, "results": [],
            "note": ("No analysis results saved yet. Run run_analysis with save_as "
                     "to remember a result for later turns."),
        }
    results = []
    for name, val in memory.items():
        item = {"name": str(name)}
        if isinstance(val, pd.DataFrame):
            item["kind"] = "dataframe"
            item["rows"] = int(len(val))
            item["columns"] = [str(c) for c in val.columns][:30]
        elif isinstance(val, pd.Series):
            item["kind"] = "series"
            item["length"] = int(len(val))
            item["name_of_series"] = None if val.name is None else str(val.name)
        elif isinstance(val, (list, tuple, set)):
            item["kind"] = "list"
            item["length"] = len(val)
        elif isinstance(val, dict):
            item["kind"] = "dict"
            item["keys"] = [str(k) for k in list(val.keys())[:30]]
        else:
            item["kind"] = "scalar"
            item["value"] = _plain(val)
        results.append(item)
    return {"count": len(results), "results": results}


# ==================== data diagnostics (read-only) ====================
#
# Thin, JSON-returning wrappers over existing dashboard analytics so the
# assistant can answer questions the cost/optimization tools don't cover:
# historical carrier share (the baseline growth caps are judged against), rate
# coverage gaps, and where specific containers are. All pure and read-only.


def _containers_sum(df) -> int:
    """Total container count of a frame, robust to a missing/garbage column."""
    if df is None or len(df) == 0 or "Container Count" not in df.columns:
        return 0
    return int(pd.to_numeric(df["Container Count"], errors="coerce").fillna(0).sum())


def historic_volume_share(df, scope_input=None, n_weeks=5, top_n=25) -> dict:
    """Tool handler: carrier market share over the last N completed weeks.

    Delegates to optimization.historic_volume. Returns, per lane (and category
    when present), each carrier's share of that lane's volume, weeks active, and
    average weekly containers — the baseline the optimizer's growth caps and any
    minimum/percent constraints are judged against. Optionally scoped.
    """
    from optimization.historic_volume import calculate_carrier_volume_share
    from config.carrier_mapping import get_carrier_name
    from .simulation import Scope
    if df is None or len(df) == 0:
        return {"error": "No data is loaded. Upload GVT and Rate files first."}
    required = ["Dray SCAC(FL)", "Container Count", "Week Number", "Lane"]
    missing = [c for c in required if c not in df.columns]
    if missing:
        return {"error": f"Data is missing column(s) needed for volume share: {missing}."}

    sub = df
    if scope_input:
        scope = Scope.from_dict(scope_input)
        if not scope.is_empty_spec():
            sub = df[scope.mask(df)]
    if sub is None or len(sub) == 0:
        return {"matched_rows": 0, "rows": [], "note": "No containers match that selection."}

    try:
        nw = max(1, int(n_weeks))
    except (TypeError, ValueError):
        nw = 5
    try:
        share = calculate_carrier_volume_share(sub, n_weeks=nw)
    except ValueError as exc:
        return {"error": f"Could not compute volume share: {exc}"}
    if share is None or len(share) == 0:
        return {"n_weeks": nw, "rows": [],
                "note": f"No historical volume found in the last {nw} week(s) for this selection."}

    has_cat = "Category" in share.columns
    try:
        cap = max(1, int(top_n))
    except (TypeError, ValueError):
        cap = 25
    rows = []
    for _, r in share.head(cap).iterrows():
        scac = str(r["Dray SCAC(FL)"])
        row = {
            "carrier": scac, "name": get_carrier_name(scac), "lane": str(r["Lane"]),
            "containers": _plain(r.get("Total_Containers")),
            "lane_total": _plain(r.get("Lane_Total_Containers")),
            "volume_share_pct": _plain(r.get("Volume_Share_Pct")),
            "weeks_active": _plain(r.get("Weeks_Active")),
            "avg_weekly_containers": _plain(r.get("Avg_Weekly_Containers")),
        }
        if has_cat:
            row["category"] = str(r.get("Category"))
        rows.append(row)
    weeks = sorted(int(w) for w in
                   pd.to_numeric(sub["Week Number"], errors="coerce").dropna().unique())[-nw:]
    return {"n_weeks": nw, "weeks_analyzed": weeks, "rows": rows,
            "rows_omitted": max(0, len(share) - cap)}


def missing_rate_audit(df, top_n=25) -> dict:
    """Tool handler: lanes/containers with a missing or non-positive rate.

    Surfaces the rate-coverage gaps the flip and optimization tools have to work
    around: how much volume can't be priced, broken down by carrier and lane.
    Mirrors components.missing_rate_analysis but returns JSON, not a Streamlit UI.
    """
    from config.carrier_mapping import get_carrier_name
    if df is None or len(df) == 0:
        return {"error": "No data is loaded. Upload GVT and Rate files first."}
    if "Base Rate" not in df.columns:
        return {"error": "Data has no Base Rate column to audit."}

    rate = pd.to_numeric(df["Base Rate"], errors="coerce")
    bad = df[rate.isna() | (rate <= 0)]
    total_containers = _containers_sum(df)
    if len(bad) == 0:
        return {"has_missing_rates": False, "total_containers": total_containers,
                "affected_containers": 0, "affected_pct": 0.0,
                "note": "All rows have a positive Base Rate; nothing is unpriced."}

    affected = _containers_sum(bad)
    pct = (affected / total_containers * 100.0) if total_containers > 0 else 0.0
    try:
        cap = max(1, int(top_n))
    except (TypeError, ValueError):
        cap = 25
    ccol = _carrier_col(bad)
    bad_c = bad.assign(_c=pd.to_numeric(bad.get("Container Count"), errors="coerce").fillna(0))

    by_lane = []
    if "Lane" in bad.columns:
        agg = bad_c.groupby(bad_c["Lane"].astype(str))["_c"].sum().sort_values(ascending=False).head(cap)
        for lane, c in agg.items():
            entry = {"lane": str(lane), "containers": int(c)}
            if ccol in bad.columns:
                entry["carriers"] = sorted({str(x).strip() for x in
                                            bad.loc[bad["Lane"].astype(str) == lane, ccol].dropna()})
            by_lane.append(entry)

    by_carrier = []
    if ccol in bad.columns:
        agg = bad_c.groupby(bad_c[ccol].astype(str).str.strip()).agg(
            containers=("_c", "sum"),
            lanes=("Lane", "nunique") if "Lane" in bad.columns else ("_c", "size"),
        ).sort_values("containers", ascending=False).head(cap)
        for scac, r in agg.iterrows():
            by_carrier.append({"carrier": str(scac), "name": get_carrier_name(str(scac)),
                               "containers": int(r["containers"]), "affected_lanes": int(r["lanes"])})

    return {
        "has_missing_rates": True, "total_containers": total_containers,
        "affected_containers": affected, "affected_pct": round(pct, 2),
        "affected_rows": int(len(bad)),
        "affected_lanes": int(bad["Lane"].nunique()) if "Lane" in bad.columns else None,
        "by_lane": by_lane, "by_carrier": by_carrier,
    }


def trace_containers(df, container_ids, max_rows=100) -> dict:
    """Tool handler: locate specific container IDs in the current data.

    For each requested ID, report the carrier currently holding it and its lane,
    week, port, facility and category. IDs not present are listed as not_found
    (never invented). Delegates parsing to components.container_tracer.
    """
    from components.reporting.container_tracer import build_container_origin_map
    from config.carrier_mapping import get_carrier_name
    if df is None or len(df) == 0:
        return {"error": "No data is loaded. Upload GVT and Rate files first."}
    if "Container Numbers" not in df.columns:
        return {"error": "Data has no 'Container Numbers' column — container-level "
                         "tracing is unavailable for this dataset."}
    if isinstance(container_ids, str):
        container_ids = [container_ids]
    wanted = [str(c).strip() for c in (container_ids or []) if str(c).strip()]
    if not wanted:
        return {"error": "No container IDs given to trace."}

    origin = build_container_origin_map(df, carrier_col=_carrier_col(df))
    upper_index = {str(k).upper(): k for k in origin}  # case-insensitive lookup
    try:
        cap = max(1, int(max_rows))
    except (TypeError, ValueError):
        cap = 100

    found, not_found = [], []
    for cid in wanted:
        key = upper_index.get(cid.upper())
        if key is None:
            not_found.append(cid)
            continue
        info = origin[key]
        scac = str(info.get("original_carrier", "")).strip()
        rec = {
            "container": cid, "carrier": scac or None,
            "carrier_name": get_carrier_name(scac) if scac else None,
            "week": _plain(info.get("week")),
            "port": info.get("discharged_port") or None,
            "lane": info.get("lane") or None,
            "facility": info.get("facility") or None,
            "category": info.get("category") or None,
        }
        if info.get("duplicate"):
            rec["note"] = (f"Appears in {info.get('duplicate_count', 2)} rows "
                           "(e.g. multiple weeks); showing the first.")
        found.append(rec)

    matched = len(wanted) - len(not_found)
    return {
        "requested": len(wanted),
        "found_count": matched,
        "found": found[:cap],
        "not_found": not_found,
        "rows_omitted": max(0, matched - cap),
    }


# ==================== constraint validation ====================

def validate_constraint(constraint: dict, valid_carriers: Optional[set] = None) -> list:
    """Return a list of human-readable problems with a single constraint dict.

    Empty list == valid. Mirrors the rules enforced by constraints_processor so
    the assistant produces constraints the optimizer will actually apply.
    """
    problems = []

    # Priority Score is the only hard-required field.
    prio = _coerce_num(constraint.get("Priority Score"))
    if prio is None:
        problems.append("Priority Score is required and must be a number.")

    has_max = _coerce_num(constraint.get("Maximum Container Count")) is not None
    has_min = _coerce_num(constraint.get("Minimum Container Count")) is not None
    has_pct = _coerce_num(constraint.get("Percent Allocation")) is not None
    has_excl = not _is_blank(constraint.get("Excluded FC"))
    carrier = constraint.get("Carrier")
    has_carrier = not _is_blank(carrier)

    # An actionable constraint needs at least one amount or an exclusion.
    if not (has_max or has_min or has_pct or has_excl):
        problems.append(
            "Constraint has no effect: set Maximum/Minimum/Percent Allocation or Excluded FC."
        )

    # Max / Min / Excluded FC all require a target carrier.
    if (has_max or has_min or has_excl) and not has_carrier:
        problems.append("A Carrier is required for Maximum, Minimum, or Excluded FC constraints.")

    # Range sanity.
    pct = _coerce_num(constraint.get("Percent Allocation"))
    if pct is not None and (pct < 0 or pct > 100):
        problems.append("Percent Allocation must be between 0 and 100.")
    mx = _coerce_num(constraint.get("Maximum Container Count"))
    if mx is not None and mx < 0:
        problems.append("Maximum Container Count cannot be negative.")
    mn = _coerce_num(constraint.get("Minimum Container Count"))
    if mn is not None and mn < 0:
        problems.append("Minimum Container Count cannot be negative.")
    if mx is not None and mn is not None and mn > mx:
        problems.append("Minimum Container Count cannot exceed Maximum Container Count.")

    wk = constraint.get("Week Number")
    if not _is_blank(wk) and _coerce_num(wk) is None:
        problems.append("Week Number must be a number if provided.")

    dow = constraint.get("Day of Week")
    if not _is_blank(dow) and parse_day_of_week(dow) is None:
        problems.append(
            "Day of Week must be a number 1-7 (1=Sunday) or a day name "
            "(mon/monday/…) if provided."
        )

    if valid_carriers is not None and has_carrier:
        if str(carrier).strip() not in valid_carriers:
            problems.append(
                f"Carrier '{carrier}' is not present in the loaded data "
                f"(it will still apply if the carrier appears in the rate card)."
            )

    return problems


def _canonical_key(key) -> str:
    """Fold a key to a separator/case-insensitive form for matching.

    Lets the model send 'Priority Score', 'priority_score', 'priorityScore',
    or 'PRIORITY SCORE' and all resolve to the same canonical column. Bedrock's
    Converse API forbids spaces in tool-schema property keys, so the schemas use
    underscores while the constraint columns use spaces — this bridges the two.
    """
    return "".join(ch for ch in str(key).lower() if ch.isalnum())


# Precomputed map: canonical(column) -> column.
_CANON_TO_COLUMN = {_canonical_key(col): col for col in CONSTRAINT_COLUMNS}


def _normalize_constraint(raw: dict) -> dict:
    """Coerce an arbitrary dict into a full constraint row with all columns."""
    row = {col: None for col in CONSTRAINT_COLUMNS}
    for k, v in raw.items():
        col = _CANON_TO_COLUMN.get(_canonical_key(k))
        if col is not None:
            row[col] = v if not _is_blank(v) else None
    # Coerce numerics.
    for col in _AMOUNT_FIELDS + ["Priority Score", "Week Number"]:
        if not _is_blank(row[col]):
            num = _coerce_num(row[col])
            row[col] = num if num is not None else row[col]
    # Parse Day of Week (number 1-7 or name like 'monday') to the Excel WEEKDAY
    # number the matcher expects. The chatbot apply path goes straight to
    # apply_constraints_to_data, bypassing process_constraints_file where this
    # parsing normally happens — so do it here too. Unrecognized values are left
    # as-is so validation can flag them rather than silently dropping the filter.
    if not _is_blank(row["Day of Week"]):
        parsed_dow = parse_day_of_week(row["Day of Week"])
        if parsed_dow is not None:
            row["Day of Week"] = parsed_dow
    # Preserve provenance tag if the source row carried one.
    if raw.get("_origin"):
        row["_origin"] = raw["_origin"]
    return row


def _working_set_summary(rows: list, df: Optional[pd.DataFrame] = None) -> list:
    """Compact, indexed view of the working set, each row tagged with scope match.

    Folds what ``describe_constraints`` and ``preview_constraint_scope`` would
    return into one structure, so a generate/edit call can hand the model the
    full picture (what's staged + how many containers each rule touches) without
    a separate round-trip. ``scope_match`` is omitted when ``df`` is unavailable.
    """
    summary = []
    for i, c in enumerate(rows):
        item = {"index": i, "origin": c.get("_origin", "assistant")}
        for col in CONSTRAINT_COLUMNS:
            v = c.get(col)
            if not _is_blank(v):
                item[col] = _plain(v)
        if c.get("_problems"):
            item["problems"] = c["_problems"]
        if df is not None:
            match = _scope_match(df, c)
            if match is not None:
                item["scope_containers"] = match["matched_containers"]
        summary.append(item)
    return summary


def generate_constraints(proposals: list, existing: Optional[list] = None,
                         valid_carriers: Optional[set] = None,
                         df: Optional[pd.DataFrame] = None) -> dict:
    """Normalize + validate proposed constraint rows, APPENDING to the existing set.

    New rows are added after any ``existing`` working-set constraints (e.g. ones
    loaded from an uploaded file) so generating a rule never silently discards
    constraints already in play. Returns the full combined working set with each
    row carrying ``_problems`` (empty if valid) and ``_origin``.

    Composite behaviour: when ``df`` (the loaded data) is supplied, the result
    also carries ``working_set`` — an indexed summary where every rule is tagged
    with how many containers its scope matches. This fuses what would otherwise
    be separate ``preview_constraint_scope`` / ``describe_constraints`` calls into
    this one tool result, so the assistant can warn "this cap won't bind" without
    an extra model round-trip.
    """
    if not isinstance(proposals, list):
        return {"error": "proposals must be a list of constraint objects."}

    # Start from the existing working set so uploaded/earlier rules are preserved.
    out_rows = [dict(c) for c in (existing or [])]
    added = 0
    for raw in proposals:
        if not isinstance(raw, dict):
            out_rows.append({"_problems": ["Constraint is not an object."], "raw": raw})
            added += 1
            continue
        row = _normalize_constraint(raw)
        row.setdefault("_origin", "assistant")
        row["_problems"] = validate_constraint(row, valid_carriers)
        out_rows.append(row)
        added += 1

    valid_count = sum(1 for r in out_rows if not r.get("_problems"))
    result = {
        "constraints": out_rows,
        "added": added,
        "valid_count": valid_count,
        "invalid_count": len(out_rows) - valid_count,
    }
    if df is not None:
        result["working_set"] = _working_set_summary(out_rows, df)
    return result


def edit_constraints(existing: list, edits: list,
                     valid_carriers: Optional[set] = None,
                     df: Optional[pd.DataFrame] = None) -> dict:
    """Apply edits to an existing constraint list.

    Each edit is a dict with an "action" ("update" | "delete" | "add") and:
      - update/delete: "index" (0-based position in `existing`)
      - update/add: the constraint fields to set.
    Returns {"constraints": [...], "applied": n, "errors": [...]}.

    Composite behaviour: when ``df`` is supplied, the result also carries an
    indexed ``working_set`` with each surviving rule's scope-match count, so the
    assistant can show the post-edit set (and whether caps bind) without a
    separate ``describe_constraints`` / ``preview_constraint_scope`` round-trip.
    """
    existing = existing or []
    result = [dict(c) for c in existing]
    errors = []
    applied = 0

    # Process deletes last (by descending index) so earlier indices stay valid.
    deletes = []
    for edit in edits:
        action = str(edit.get("action", "update")).lower()
        if action == "add":
            row = _normalize_constraint(edit)
            row["_problems"] = validate_constraint(row, valid_carriers)
            result.append(row)
            applied += 1
        elif action == "update":
            idx = edit.get("index")
            if not isinstance(idx, int) or idx < 0 or idx >= len(result):
                errors.append(f"update: index {idx} out of range (0..{len(result)-1}).")
                continue
            merged = dict(result[idx])
            for k, v in edit.items():
                if k in ("action", "index"):
                    continue
                col = _CANON_TO_COLUMN.get(_canonical_key(k))
                if col is not None:
                    merged[col] = v if not _is_blank(v) else None
            merged = _normalize_constraint(merged)
            merged["_problems"] = validate_constraint(merged, valid_carriers)
            result[idx] = merged
            applied += 1
        elif action == "delete":
            idx = edit.get("index")
            if not isinstance(idx, int) or idx < 0 or idx >= len(result):
                errors.append(f"delete: index {idx} out of range (0..{len(result)-1}).")
                continue
            deletes.append(idx)
        else:
            errors.append(f"Unknown action '{action}'.")

    for idx in sorted(set(deletes), reverse=True):
        result.pop(idx)
        applied += 1

    out = {"constraints": result, "applied": applied, "errors": errors}
    if df is not None:
        out["working_set"] = _working_set_summary(result, df)
    return out


# ==================== scope preview (how many containers would a constraint hit) ====================

def _scope_match(df: Optional[pd.DataFrame], constraint: dict) -> Optional[dict]:
    """Core scope-match counter shared by preview and the composite constraint tools.

    Returns {matched_containers, matched_rows, filters_applied} for a constraint's
    scope filters (Carrier is ignored — it's the assignment target, not a filter),
    or None if the data can't support a count (no df / no Container Count column).
    Keys may be underscored ('week_number') or spaced ('Week Number'); both resolve.
    """
    if df is None or len(df) == 0 or "Container Count" not in df.columns:
        return None

    # Resolve incoming keys (which may use underscores, e.g. 'week_number')
    # to canonical column names before filtering.
    resolved = {}
    for k, v in (constraint or {}).items():
        col = _CANON_TO_COLUMN.get(_canonical_key(k))
        if col is not None:
            resolved[col] = v

    mask = pd.Series(True, index=df.index)
    applied_filters = {}
    for field in _SCOPE_FIELDS:
        if field == "Carrier":
            # Carrier is the assignment target, not a scope filter — skip.
            continue
        val = resolved.get(field)
        if _is_blank(val):
            continue
        col = _DATA_COLUMN_FOR.get(field)
        if col not in df.columns:
            continue
        if field == "Week Number":
            num = _coerce_num(val)
            mask &= df[col] == num
        elif field == "Day of Week":
            # Data carries the Excel WEEKDAY number; parse the value (name or number)
            # so the preview matches even if it wasn't pre-normalized.
            dow = parse_day_of_week(val)
            mask &= pd.to_numeric(df[col], errors="coerce") == dow
        elif field == "Lane":
            lane_val = str(val).strip()
            if len(lane_val) <= 4:
                mask &= df[col].astype(str).str.endswith(lane_val)
            else:
                mask &= df[col].astype(str) == lane_val
        else:
            mask &= df[col].astype(str).str.strip() == str(val).strip()
        applied_filters[field] = val

    matched = df[mask]
    matched_containers = int(matched["Container Count"].sum()) if len(matched) else 0
    return {
        "matched_containers": matched_containers,
        "matched_rows": int(len(matched)),
        "filters_applied": applied_filters or "none (matches all data)",
    }


def preview_constraint_scope(df: Optional[pd.DataFrame], constraint: dict) -> dict:
    """Count how many containers in the current data a constraint's scope matches.

    Helps the assistant (and user) sanity-check a proposed rule before applying.
    Note: this is a scope preview, not the full allocation the optimizer runs.
    """
    if df is None or len(df) == 0:
        return {"error": "No data loaded."}
    if "Container Count" not in df.columns:
        return {"error": "Container Count column missing."}
    return _scope_match(df, constraint)


# ==================== working-set inspection ====================

def describe_constraints(constraints: Optional[list], source: Optional[str] = None) -> dict:
    """Report the current working set of constraints for the assistant.

    Returns each constraint with its 0-based index (so the assistant can target
    it with edit_constraints), its origin ('uploaded' vs 'assistant'), and any
    validation problems. When the set is empty, returns a note telling the
    assistant to ask the user to upload a file or describe a rule — it must NOT
    invent constraints to edit.
    """
    constraints = constraints or []
    if not constraints:
        return {
            "count": 0,
            "source": source,
            "constraints": [],
            "note": (
                "There are no constraints loaded. Do not invent any. Ask the user "
                "to either upload a constraint file (the main uploader) or describe "
                "the rule they want — then use generate_constraints."
            ),
        }
    rows = []
    for i, c in enumerate(constraints):
        row = {"index": i, "origin": c.get("_origin", "assistant")}
        for col in CONSTRAINT_COLUMNS:
            v = c.get(col)
            if not _is_blank(v):
                row[col] = _plain(v)
        if c.get("_problems"):
            row["problems"] = c["_problems"]
        rows.append(row)
    return {"count": len(rows), "source": source, "constraints": rows}


# ==================== applied-constraint outcome ====================
#
# The dashboard's "Applied Constraints Summary" (built by
# constraints_processor.apply_constraints_to_data and rendered by
# show_constraints_summary) records, per rule, how much volume it actually
# claimed vs targeted, which higher-priority rules consumed its pool, and why a
# minimum/percent rule fell short. The dashboard stashes that list into session
# state each run; this tool digests it so the assistant can explain the real
# impact of the live constraints and ground new-constraint suggestions in what
# the optimizer actually did — not just what the rules say on paper.


def _format_claimed_by(claimed_by) -> Optional[dict]:
    """Render the {priority: containers} claim map into JSON-safe 'P<n>' keys."""
    if not claimed_by:
        return None
    return {
        f"P{_plain(prio)}": _plain(count)
        for prio, count in claimed_by.items()
    }


# Scope dimensions that partition the *run's* data. When one of these is absent
# from the GVT, it usually means this run simply doesn't cover that segment (e.g.
# a CD rule on a Robotics/Devices-only run) — a failure that is EXPECTED, not a
# misconfiguration. The matching rule will bind the moment that volume is loaded.
_COARSE_PARTITION_DIMS = {"Category", "Port"}
# Fine-grained dimensions. A value here being absent from the data is more likely
# a typo, a stale code, or a value that lives under a different week/port — worth
# a human look regardless of priority.
_FINE_GRAINED_DIMS = {"Lane", "Terminal", "Vessel", "SSL", "Week", "Week Number", "Day of Week"}

# Failure classes. The first two are "fine to fail" (the rule is well-formed; the
# run just can't honour it); the rest are genuine misconfigurations to surface.
_FAIL_OUT_OF_SCOPE = "out_of_scope_data"      # acceptable
_FAIL_SUPERSEDED = "superseded"               # acceptable
_FAIL_DEAD_FILTER = "dead_filter_value"       # needs attention
_FAIL_NARROW_COMBO = "narrow_combination"     # needs attention
_FAIL_EXCLUSION = "exclusion_conflict"        # needs attention
_FAIL_MALFORMED = "malformed"                 # needs attention
_FAIL_UNKNOWN = "unclassified"                # needs attention (be conservative)

_ACCEPTABLE_FAIL_CLASSES = {_FAIL_OUT_OF_SCOPE, _FAIL_SUPERSEDED}


def _dead_dimensions_from_reason(reason: str) -> list:
    """Pull the dead scope dimension names out of a 'No matching data' reason.

    The processor writes reasons like '...scope filter(s) Category=CD, Lane=RMN3.
    That value isn't present...'. Returns ['Category', 'Lane'] for that example,
    or [] when the reason isn't a dead-value diagnosis.
    """
    if not reason:
        return []
    m = re.search(r"scope filter\(s\)\s+(.+?)\.\s+That value", reason)
    if not m:
        return []
    dims = []
    for part in m.group(1).split(","):
        part = part.strip()
        if "=" in part:
            dims.append(part.split("=", 1)[0].strip())
    return dims


def classify_constraint_failure(status: str, reason: Optional[str],
                                claimed_by_present: bool) -> dict:
    """Classify WHY a constraint produced no allocation — root cause first.

    Returns ``{"class", "acceptable", "rationale"}``. Per the product rule, the
    *root cause* decides whether a failure is fine to ignore (priority is only
    used later to rank the ones that do need attention):

      - superseded / out_of_scope_data -> acceptable at ANY priority. A higher-
        priority rule legitimately claimed the volume, or the run just doesn't
        contain that segment. Nothing to fix.
      - dead_filter_value / narrow_combination / exclusion_conflict / malformed
        -> needs attention at ANY priority (a typo'd lane stays a typo whether
        it's P10 or P1).

    Only meaningful for non-allocating rules. Applied/Partial rules return the
    sentinel class ``""`` with acceptable=True (they are not failures).
    """
    status = str(status or "")
    reason = reason or ""

    # Not a failure: applied (incl. lockouts/exclusion rules) or a partial fill.
    if status.startswith(("Applied", "Partial")):
        return {"class": "", "acceptable": True, "rationale": ""}

    # Malformed rows (missing carrier, allocated-to-excluded-facility rollback).
    if status.startswith(("Error", "FAILED")):
        return {
            "class": _FAIL_MALFORMED, "acceptable": False,
            "rationale": "The rule itself is malformed or contradictory — fix the row.",
        }

    # Superseded: a higher-priority rule claimed the contested containers first.
    # This is the priority hierarchy working as designed.
    if claimed_by_present or "already claimed by higher-priority" in reason:
        return {
            "class": _FAIL_SUPERSEDED, "acceptable": True,
            "rationale": ("A higher-priority rule claimed this scope's volume first — "
                          "expected behaviour. Only act if THIS rule should outrank it."),
        }

    # Exclusion wiped the pool: scope + excluded-facility filter left nothing.
    if "excluded facilit" in reason and "after removing" in reason:
        return {
            "class": _FAIL_EXCLUSION, "acceptable": False,
            "rationale": ("Scope plus the facility exclusion eliminated every container. "
                          "Widen the scope or relax the exclusion — but first confirm the "
                          "scope's Category/Port even exists in this run (it may really be "
                          "out-of-scope data, which the facility message can mask)."),
        }

    # Dead value: a scope filter matched zero rows in the source data.
    dead_dims = _dead_dimensions_from_reason(reason)
    if dead_dims or "isn't present in the GVT" in reason:
        # A coarse partition (Category/Port) being absent means the run doesn't
        # cover that segment -> expected. Only when the dead value is fine-grained
        # (a lane/terminal/vessel/etc.) is it a likely typo worth flagging.
        coarse_dead = [d for d in dead_dims if d in _COARSE_PARTITION_DIMS]
        if coarse_dead:
            return {
                "class": _FAIL_OUT_OF_SCOPE, "acceptable": True,
                "rationale": (f"This run's data has no {', '.join(coarse_dead)} matching the "
                              "rule, so it can't bind here — fine to fail if that segment "
                              "isn't in scope this run; it'll apply once that volume loads."),
            }
        return {
            "class": _FAIL_DEAD_FILTER, "acceptable": False,
            "rationale": (f"A {', '.join(dead_dims) or 'scope'} value isn't in the data — "
                          "likely a typo, a stale code, or a value under a different "
                          "week/port. Correct it or drop that filter."),
        }

    # Filter combination too narrow (each filter matches alone, none together).
    if "no single row satisfies all of them" in reason:
        return {
            "class": _FAIL_NARROW_COMBO, "acceptable": False,
            "rationale": ("Each filter matches on its own but their AND-combination is "
                          "empty — relax one dimension so the scopes overlap."),
        }

    # Be conservative with anything we can't recognise: flag for review.
    return {
        "class": _FAIL_UNKNOWN, "acceptable": False,
        "rationale": "Produced no allocation for an unrecognised reason — review manually.",
    }


def summarize_applied_constraints(constraint_summary: Optional[list]) -> dict:
    """Digest the dashboard's Applied Constraints Summary for the assistant.

    `constraint_summary` is the list of per-rule outcome dicts the dashboard
    produces (priority, status, containers_allocated, target_containers,
    eligible_containers, claimed_by, scope, reason, ...). Returns a JSON-safe
    rollup plus a per-rule breakdown and an actionable ``shortfalls`` list of
    minimum/percent rules that could not be fully satisfied — the natural seed
    for suggesting new or adjusted constraints.

    Returns {"applied": False, "note": ...} when no constraints were applied so
    the assistant tells the user to apply constraints rather than inventing an
    impact report.
    """
    if not constraint_summary or not isinstance(constraint_summary, list):
        return {
            "applied": False,
            "note": (
                "No constraints have been applied to the current optimization, so "
                "there is no Applied Constraints Summary to read. Ask the user to "
                "upload a constraint file or draft and Apply constraints first; do "
                "not fabricate an impact report."
            ),
        }

    rules = []
    shortfalls = []
    acceptable_failures = []   # well-formed rules the run can't honour (fine to fail)
    needs_attention = []       # genuine misconfigurations to surface to the user
    total_allocated = 0
    successful = partial = failed_or_skipped = 0

    for item in constraint_summary:
        if not isinstance(item, dict):
            continue
        status = str(item.get("status", "")) or "Unknown"
        allocated = _plain(item.get("containers_allocated", 0)) or 0
        target = item.get("target_containers")
        target = _plain(target) if target is not None else None
        eligible = item.get("eligible_containers")

        if isinstance(allocated, (int, float)):
            total_allocated += allocated
        if status.startswith("Applied"):
            successful += 1
        elif status.startswith("Partial"):
            partial += 1
        else:
            failed_or_skipped += 1

        rule = {
            "priority": _plain(item.get("priority")),
            "status": status,
            "allocated": allocated,
        }
        if item.get("description"):
            rule["description"] = item["description"]
        if target is not None:
            rule["target"] = target
        if eligible is not None:
            rule["eligible"] = _plain(eligible)
        if item.get("method"):
            rule["method"] = item["method"]
        scope = item.get("scope")
        if scope:
            rule["scope"] = {k: _plain_scope_val(v) for k, v in scope.items()}
        claimed = _format_claimed_by(item.get("claimed_by"))
        if claimed:
            rule["claimed_by"] = claimed
        if item.get("reason"):
            rule["why"] = item["reason"]

        # Root-cause triage for non-allocating rules. Whether a failure is "fine"
        # is decided by WHY it failed, not its priority; priority only ranks the
        # ones that do need attention (see classify_constraint_failure).
        is_nonalloc_fail = (
            not status.startswith(("Applied", "Partial"))
            or (status.startswith("Partial") and allocated == 0)
        )
        if is_nonalloc_fail:
            verdict = classify_constraint_failure(
                status, item.get("reason"), bool(claimed)
            )
            rule["failure_class"] = verdict["class"]
            rule["acceptable_to_fail"] = verdict["acceptable"]
            rule["triage_note"] = verdict["rationale"]
            bucket = acceptable_failures if verdict["acceptable"] else needs_attention
            bucket.append(rule)
        rules.append(rule)

        # An actionable shortfall: a rule that wanted volume but missed its
        # target (partial fill, or a failed/errored min/percent rule).
        is_short = status.startswith(("Partial", "Failed", "FAILED", "Error"))
        if not is_short and isinstance(target, (int, float)) and target > 0 \
                and isinstance(allocated, (int, float)) and allocated < target:
            is_short = True
        if is_short and isinstance(target, (int, float)) and target > 0:
            short = {
                "priority": rule["priority"],
                "status": status,
                "target": target,
                "allocated": allocated,
                "gap": _plain(max(0, target - allocated))
                if isinstance(allocated, (int, float)) else target,
            }
            if rule.get("scope"):
                short["scope"] = rule["scope"]
            if claimed:
                short["claimed_by"] = claimed
            if item.get("reason"):
                short["why"] = item["reason"]
            shortfalls.append(short)

    # Rank the genuine problems so the assistant leads with the most important.
    # Priority is the tiebreak HERE (within already-flagged rules), never the
    # gate for whether a failure counts (root cause decides that, above).
    def _prio_key(r):
        p = r.get("priority")
        return -p if isinstance(p, (int, float)) else 0
    needs_attention.sort(key=_prio_key)

    failure_classes = {}
    for r in acceptable_failures + needs_attention:
        cls = r.get("failure_class") or _FAIL_UNKNOWN
        failure_classes[cls] = failure_classes.get(cls, 0) + 1

    return {
        "applied": True,
        "rule_count": len(rules),
        "total_allocated_containers": _plain(total_allocated),
        "successful": successful,
        "partial": partial,
        "failed_or_skipped": failed_or_skipped,
        # Triage of the failures: acceptable_to_fail (well-formed rules the run
        # can't honour — superseded by higher priority, or out-of-scope data) vs.
        # needs_attention (typos, over-narrow scope, exclusion conflicts,
        # malformed rows). LEAD with needs_attention; treat acceptable as FYI.
        "acceptable_failure_count": len(acceptable_failures),
        "needs_attention_count": len(needs_attention),
        "failure_classes": failure_classes,
        "acceptable_failures": acceptable_failures,
        "needs_attention": needs_attention,
        "rules": rules,
        "shortfalls": shortfalls,
    }


# ==================== deep constraint analysis & repair ====================
#
# diagnose_constraints/repair_constraints turn the working set + the loaded data
# (and, when present, the applied-constraints outcome) into a structured report
# of the THREE failure modes a hand-written constraint list keeps hitting:
#   1. over-subscription — fixed caps + percent rules for one (Port, Category)
#      scope together ask for more containers than the scope holds. The processor
#      computes each percent against the FROZEN ORIGINAL pool, so summed percents
#      can exceed 100% of what's actually left; lower-priority rules then starve.
#   2. tiny pools — a scope with very few containers but many rules; the
#      lower-priority rules can never bind because higher ones drain the handful.
#   3. dead scopes — a rule whose scope matches zero rows. Split (via the same
#      root-cause triage as summarize_applied_constraints) into fixable typos
#      (dead_filter_value) vs acceptable out-of-scope-this-run rules.
# repair_constraints then produces a corrected working set: rescale over-subscribed
# percents to fit the pool, drop only the fixable dead rules, and collapse
# redundant rules on a tiny pool to the carriers actually present there. Lockouts
# (0% / max 0) are NEVER rescaled or dropped — they are meaningful blocks.

# Thresholds for the "tiny pool, many rules" heuristic. A scope at or below
# TINY_POOL_CONTAINERS containers carrying at least TINY_POOL_MIN_RULES rules is
# over-ruled: the volume can't satisfy that many competing rules.
TINY_POOL_CONTAINERS = 10
TINY_POOL_MIN_RULES = 3


def _canonical_scope_key(constraint: dict):
    """(port_bucket, category_bucket) a constraint groups under for pool math.

    Port is upper-cased (alias expansion is handled by the data side); Category is
    routed through the shared canonicalizer so 'Retail CD'/'FBA FCL'/'CD' all
    collapse to one bucket. Either side may be None (means 'all' for that
    dimension), and a (None, None) scope is the whole run.
    """
    from config.category_mapping import canonical_category
    port = constraint.get("Port")
    cat = constraint.get("Category")
    port_key = str(port).strip().upper() if not _is_blank(port) else None
    cat_key = canonical_category(cat) if not _is_blank(cat) else None
    return (port_key, cat_key)


def _rule_label(constraint: dict, index: Optional[int] = None) -> dict:
    """Compact, JSON-safe identity of a rule for a diagnosis payload."""
    label = {}
    if index is not None:
        label["index"] = index
    for col in ("Priority Score", "Carrier", "Port", "Category", "Lane",
                "Week Number", "Day of Week", "Maximum Container Count",
                "Minimum Container Count", "Percent Allocation", "Excluded FC"):
        v = constraint.get(col)
        if not _is_blank(v):
            label[col] = _plain(v)
    return label


def _is_lockout(constraint: dict) -> bool:
    """True if the rule is a lockout (0% or Max 0) — a meaningful block, not a no-op."""
    pct = _coerce_num(constraint.get("Percent Allocation"))
    mx = _coerce_num(constraint.get("Maximum Container Count"))
    return pct == 0 or mx == 0


def diagnose_constraints(constraints: Optional[list], df: Optional[pd.DataFrame],
                         constraint_summary: Optional[list] = None) -> dict:
    """Deep-analyze a constraint working set against the loaded data.

    Detects over-subscribed scopes (caps + percents exceed the pool), tiny pools
    carrying too many rules, and dead scopes (zero matching containers, split into
    fixable typos vs acceptable out-of-scope rules). When an applied-constraints
    ``constraint_summary`` is supplied, its shortfalls/needs_attention are folded
    in so the report reflects what the optimizer actually did, not just the rules
    on paper. Returns a JSON-safe structured report plus a ``recommended_fixes``
    list that repair_constraints consumes.
    """
    constraints = [c for c in (constraints or []) if isinstance(c, dict)]
    if not constraints:
        return {
            "analyzable": False,
            "note": ("No constraints in the working set to diagnose. Ask the user to "
                     "upload a constraint file or draft rules first."),
        }
    have_data = df is not None and len(df) > 0 and "Container Count" in getattr(df, "columns", [])

    # ---- group rules by (Port, Category) scope and size each scope's pool ----
    scopes: dict = {}
    for i, c in enumerate(constraints):
        key = _canonical_scope_key(c)
        scopes.setdefault(key, []).append((i, c))

    scope_reports = []
    over_subscribed = []
    for (port_key, cat_key), members in scopes.items():
        # Pool = containers matching just the (Port, Category) scope (no carrier).
        pool = None
        if have_data:
            probe = {}
            if port_key is not None:
                probe["Port"] = port_key
            if cat_key is not None:
                probe["Category"] = cat_key
            match = _scope_match(df, probe)
            pool = match["matched_containers"] if match else None

        # Sum the demand this scope's rules place on the pool.
        fixed_demand = 0          # Max caps + Min floors (absolute counts)
        pct_demand = 0.0          # percent share of the pool
        pct_rules, fixed_rules = [], []
        for idx, c in members:
            if _is_lockout(c):
                continue  # lockouts don't consume the pool, they block a carrier
            mx = _coerce_num(c.get("Maximum Container Count"))
            mn = _coerce_num(c.get("Minimum Container Count"))
            pct = _coerce_num(c.get("Percent Allocation"))
            # A lane-scoped rule draws from a sub-pool, not the whole (Port,Cat)
            # pool, so it isn't additive here — note it but don't sum it.
            lane_scoped = not _is_blank(c.get("Lane"))
            if pct is not None and pct > 0 and not lane_scoped:
                pct_demand += pct
                pct_rules.append((idx, c, pct))
            if (mx is not None and mx > 0) or (mn is not None and mn > 0):
                amt = mx if (mx is not None and mx > 0) else mn
                if not lane_scoped:
                    fixed_demand += amt
                    fixed_rules.append((idx, c, amt))

        requested = None
        is_over = False
        if pool is not None:
            requested = int(round(fixed_demand + pool * (pct_demand / 100.0)))
            is_over = requested > pool and (pct_demand > 0 or fixed_demand > 0)

        report = {
            "port": port_key, "category": cat_key,
            "available_containers": pool,
            "percent_requested": round(pct_demand, 2),
            "fixed_requested": int(fixed_demand),
            "total_requested_containers": requested,
            "rule_count": len(members),
            "over_subscribed": is_over,
            "rules": [_rule_label(c, idx) for idx, c in members],
        }
        scope_reports.append(report)
        if is_over:
            over_subscribed.append(report)

    # ---- per-rule scope-match: dead scopes + tiny-pool detection ----
    dead_fixable, dead_acceptable, tiny_pools = [], [], []
    rule_scope_containers = {}
    if have_data:
        for i, c in enumerate(constraints):
            match = _scope_match(df, c)
            cnt = match["matched_containers"] if match else None
            rule_scope_containers[i] = cnt
            if cnt == 0:
                # Reuse the root-cause triage: a fine-grained dead value (lane/
                # terminal/...) is a fixable typo; a coarse absent Category/Port is
                # acceptable (this run just lacks that segment).
                dims = [d for d in ("Lane", "Terminal", "Vessel", "SSL", "Week Number", "Day of Week")
                        if not _is_blank(c.get(d))]
                coarse_only = not dims and (
                    not _is_blank(c.get("Category")) or not _is_blank(c.get("Port")))
                entry = _rule_label(c, i)
                entry["dead_dimensions"] = dims or (["Category/Port"] if coarse_only else [])
                if dims:
                    dead_fixable.append(entry)
                else:
                    dead_acceptable.append(entry)

        # Tiny pool: a (Port, Category) scope with a small pool but many rules.
        for rep in scope_reports:
            pool = rep["available_containers"]
            if pool is not None and pool <= TINY_POOL_CONTAINERS \
                    and rep["rule_count"] >= TINY_POOL_MIN_RULES:
                carriers_present = []
                if rep["port"] is not None or rep["category"] is not None:
                    probe = {}
                    if rep["port"] is not None:
                        probe["Port"] = rep["port"]
                    if rep["category"] is not None:
                        probe["Category"] = rep["category"]
                    carriers_present = _carriers_in_scope(df, probe)
                tiny_pools.append({
                    "port": rep["port"], "category": rep["category"],
                    "available_containers": pool, "rule_count": rep["rule_count"],
                    "carriers_present": carriers_present,
                    "rules": rep["rules"],
                })

    # ---- fold in the applied-outcome triage if a summary is available ----
    applied = None
    if constraint_summary:
        applied = summarize_applied_constraints(constraint_summary)

    # ---- recommended fixes (the seed repair_constraints consumes) ----
    fixes = []
    for rep in over_subscribed:
        fixes.append({
            "type": "rescale_oversubscription",
            "port": rep["port"], "category": rep["category"],
            "detail": (f"{rep['rule_count']} rule(s) request "
                       f"{rep['total_requested_containers']} containers from a "
                       f"{rep['available_containers']}-container pool. Rescale percent "
                       f"rules so caps + percents fit the pool."),
        })
    for tp in tiny_pools:
        fixes.append({
            "type": "collapse_tiny_pool",
            "port": tp["port"], "category": tp["category"],
            "detail": (f"{tp['rule_count']} rules target a {tp['available_containers']}-"
                       f"container pool (carriers present: {tp['carriers_present'] or 'none'}). "
                       f"Collapse to the viable carrier(s)."),
        })
    if dead_fixable:
        fixes.append({
            "type": "drop_dead_rules",
            "detail": (f"{len(dead_fixable)} rule(s) target a scope value not in the "
                       f"data (likely typos / stale codes). Drop or correct them."),
            "rules": dead_fixable,
        })

    return {
        "analyzable": True,
        "have_data": bool(have_data),
        "rule_count": len(constraints),
        "scopes": scope_reports,
        "over_subscribed_scopes": over_subscribed,
        "tiny_pools": tiny_pools,
        "dead_scopes": {"fixable": dead_fixable, "acceptable": dead_acceptable},
        "applied_outcome": applied,
        "recommended_fixes": fixes,
        "note": (
            "over_subscribed_scopes/tiny_pools/dead_scopes.fixable are the real "
            "problems; dead_scopes.acceptable are rules whose segment just isn't in "
            "this run (fine to leave). Use repair_constraints to stage a corrected set."
        ),
    }


def _carriers_in_scope(df: Optional[pd.DataFrame], probe: dict) -> list:
    """Carrier SCACs that actually have containers in a (Port, Category) scope."""
    if df is None or len(df) == 0:
        return []
    carrier_col = _carrier_col(df)
    if carrier_col not in df.columns:
        return []
    mask = pd.Series(True, index=df.index)
    if not _is_blank(probe.get("Port")) and "Discharged Port" in df.columns:
        mask &= df["Discharged Port"].astype(str).str.strip().str.upper() == str(probe["Port"]).strip().upper()
    if not _is_blank(probe.get("Category")) and "Category" in df.columns:
        from config.category_mapping import canonical_category
        want = canonical_category(probe["Category"])
        mask &= df["Category"].map(lambda v: canonical_category(v) == want)
    sub = df[mask]
    if not len(sub):
        return []
    return sorted({str(c).strip() for c in sub[carrier_col].dropna().unique()})


def repair_constraints(constraints: Optional[list], df: Optional[pd.DataFrame],
                       constraint_summary: Optional[list] = None,
                       valid_carriers: Optional[set] = None) -> dict:
    """Produce a corrected constraint working set (full cleanup), staged for review.

    Consumes diagnose_constraints and applies three fixes:
      * rescale over-subscribed percent rules so fixed caps + percents fit the pool;
      * drop only fixable dead-scope rules (typos), keeping out-of-scope + lockouts;
      * collapse redundant rules on a tiny pool to the carriers actually present.
    Every surviving row is re-normalized + re-validated, stays 0-100 for percents,
    and round-trips through the real processor. NEVER auto-applies — the caller
    stages the result for the user to review. Returns
    {"constraints": [...], "changes": [...], "summary": {...}}.
    """
    constraints = [dict(c) for c in (constraints or []) if isinstance(c, dict)]
    if not constraints:
        return {"constraints": [], "changes": [],
                "summary": {"rescaled": 0, "dropped": 0, "collapsed": 0,
                            "clamped": 0, "kept": 0, "original_count": 0},
                "note": "Nothing to repair: the working set is empty."}

    diag = diagnose_constraints(constraints, df, constraint_summary)
    changes = []

    # Index the dead-fixable rules and tiny-pool scopes for quick lookup.
    dead_idx = {r["index"] for r in diag.get("dead_scopes", {}).get("fixable", [])
                if "index" in r}
    over_scopes = {(s["port"], s["category"]): s
                   for s in diag.get("over_subscribed_scopes", [])}
    tiny_scopes = {(t["port"], t["category"]): t for t in diag.get("tiny_pools", [])}

    # ---- 1. collapse tiny-pool rules: keep lockouts + the highest-priority rule
    #         per carrier-present; drop rules for carriers absent from the pool.
    collapse_drop = set()
    collapsed_n = 0
    for (port_key, cat_key), tp in tiny_scopes.items():
        present = {str(c).upper() for c in tp.get("carriers_present", [])}
        for i, c in enumerate(constraints):
            if _canonical_scope_key(c) != (port_key, cat_key):
                continue
            if _is_lockout(c):
                continue  # keep blocks
            carrier = str(c.get("Carrier") or "").strip().upper()
            if carrier and present and carrier not in present and _is_blank(c.get("Lane")):
                collapse_drop.add(i)
                collapsed_n += 1
                changes.append({
                    "action": "drop_tiny_pool_rule", "index": i,
                    "scope": {"port": port_key, "category": cat_key},
                    "before": _rule_label(c, i),
                    "why": (f"Carrier {carrier} has no containers in this "
                            f"{tp['available_containers']}-container pool."),
                })

    # ---- 2. drop fixable dead-scope rules (typos / stale codes) ----
    # A LOCKOUT (0% / Max 0) is the user's explicit "block this carrier here"
    # intent — keep it even when its scope value is absent from THIS run's data
    # (it will bind once that volume loads, same rationale as out_of_scope_data).
    # Only non-lockout dead rules are dropped.
    dropped_n = 0
    dead_drop = set()
    for i in dead_idx:
        if i in collapse_drop or _is_lockout(constraints[i]):
            continue
        dead_drop.add(i)
        changes.append({
            "action": "drop_dead_rule", "index": i,
            "before": _rule_label(constraints[i], i),
            "why": "Scope value is not present in the data (likely a typo or stale code).",
        })
        dropped_n += 1
    drop_all = collapse_drop | dead_drop

    # ---- 3. rescale over-subscribed percent rules to fit the pool ----
    rescaled_n = 0
    rescale_factor = {}  # (port,cat) -> factor to multiply surviving percents by
    for (port_key, cat_key), s in over_scopes.items():
        pool = s.get("available_containers")
        if not pool or pool <= 0:
            continue
        # Re-sum the SURVIVING (not-dropped, non-lockout, non-lane) demand.
        fixed = 0.0
        pct_total = 0.0
        for i, c in enumerate(constraints):
            if i in drop_all or _is_lockout(c) or not _is_blank(c.get("Lane")):
                continue
            if _canonical_scope_key(c) != (port_key, cat_key):
                continue
            mx = _coerce_num(c.get("Maximum Container Count"))
            mn = _coerce_num(c.get("Minimum Container Count"))
            pct = _coerce_num(c.get("Percent Allocation"))
            if mx and mx > 0:
                fixed += mx
            elif mn and mn > 0:
                fixed += mn
            elif pct and pct > 0:
                pct_total += pct
        # Percent budget left after fixed caps claim their share of the pool.
        pct_budget = max(0.0, 100.0 - (fixed / pool * 100.0))
        if pct_total > pct_budget and pct_total > 0:
            rescale_factor[(port_key, cat_key)] = pct_budget / pct_total

    out_rows = []
    for i, c in enumerate(constraints):
        if i in drop_all:
            continue
        row = dict(c)
        key = _canonical_scope_key(c)
        factor = rescale_factor.get(key)
        if factor is not None and not _is_lockout(c) and _is_blank(c.get("Lane")):
            pct = _coerce_num(c.get("Percent Allocation"))
            if pct and pct > 0:
                new_pct = round(pct * factor, 1)
                changes.append({
                    "action": "rescale_percent", "index": i,
                    "scope": {"port": key[0], "category": key[1]},
                    "before": {"Percent Allocation": _plain(pct)},
                    "after": {"Percent Allocation": new_pct},
                    "why": ("Percent rules in this scope summed above the available "
                            "pool; rescaled proportionally to fit."),
                })
                row["Percent Allocation"] = new_pct
                rescaled_n += 1
        out_rows.append((i, row))

    # ---- 4. clamp out-of-range amounts so the corrected set is genuinely VALID,
    #         not merely flagged. The point of "repair" is to hand back a set the
    #         optimizer accepts; a surviving -10% or 130% (which the rescale pass
    #         skips because the row is negative, lane-scoped, or on a dead scope)
    #         would otherwise still carry a validation problem. Clamp percents to
    #         [0,100] and negative caps to 0, recording each as a change. Lockouts
    #         (already 0) are unaffected.
    clamped_n = 0
    for idx, row in out_rows:
        pct = _coerce_num(row.get("Percent Allocation"))
        if pct is not None and (pct < 0 or pct > 100):
            new_pct = 0.0 if pct < 0 else 100.0
            changes.append({
                "action": "clamp_percent", "index": idx,
                "before": {"Percent Allocation": _plain(pct)},
                "after": {"Percent Allocation": new_pct},
                "why": "Percent Allocation was outside 0-100; clamped to a valid value.",
            })
            row["Percent Allocation"] = new_pct
            clamped_n += 1
        for amt_col in ("Maximum Container Count", "Minimum Container Count"):
            amt = _coerce_num(row.get(amt_col))
            if amt is not None and amt < 0:
                changes.append({
                    "action": "clamp_amount", "index": idx,
                    "before": {amt_col: _plain(amt)},
                    "after": {amt_col: 0},
                    "why": f"{amt_col} was negative; clamped to 0.",
                })
                row[amt_col] = 0
                clamped_n += 1

    # Re-normalize + re-validate every surviving row so the corrected set is
    # guaranteed processor-acceptable; preserve origin, tag the repaired ones.
    repaired = []
    for i, row in out_rows:
        norm = _normalize_constraint(row)
        norm["_origin"] = row.get("_origin", "assistant")
        norm["_problems"] = validate_constraint(norm, valid_carriers)
        repaired.append(norm)

    result = {
        "constraints": repaired,
        "changes": changes,
        "summary": {
            "rescaled": rescaled_n,
            "dropped": dropped_n,
            "collapsed": collapsed_n,
            "clamped": clamped_n,
            "kept": len(repaired),
            "original_count": len(constraints),
        },
        "note": ("Corrected set is STAGED for review — not applied. Lockouts and "
                 "out-of-scope rules were preserved. Review in the panel, then Apply "
                 "or Download."),
    }
    if df is not None:
        result["working_set"] = _working_set_summary(repaired, df)
    return result


# ==================== export ====================

def constraints_to_dataframe(constraints: list) -> pd.DataFrame:
    """Build a clean constraint DataFrame (template column order, no helper keys)."""
    rows = []
    for c in constraints or []:
        rows.append({col: c.get(col) for col in CONSTRAINT_COLUMNS})
    df = pd.DataFrame(rows, columns=CONSTRAINT_COLUMNS)
    return df


def constraints_to_excel_bytes(constraints: list) -> bytes:
    """Serialize constraints to an .xlsx matching docs/constraint_template.xlsx."""
    df = constraints_to_dataframe(constraints)
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Constraints")
    buf.seek(0)
    return buf.getvalue()


# ==================== analysis report artifacts (xlsx + docx) ====================
#
# Pure byte-builders that turn a diagnose_constraints report (and, optionally, a
# repair_constraints corrected set) into downloadable artifacts. Both are pure:
# they take plain dicts/lists and return bytes, so chat_ui can stash the bytes in
# session_state and render a download button without these functions touching
# Streamlit. openpyxl is a hard dependency; python-docx is guarded (the Word
# builder returns None if it isn't installed, so a clean deploy degrades to xlsx).

_REPORT_NAVY = "1F4E78"
_REPORT_STATUS_FILL = {
    "over": "FFC7CE", "tiny": "FFEB9C", "dead": "FFC7CE",
    "ok": "C6EFCE", "info": "D9D9D9",
}


def _scope_label(port, category) -> str:
    return _sanitize_cell(f"{port or 'ALL'} / {category or 'ALL'}")


# Control chars openpyxl/lxml reject on write (C0 minus tab/newline/CR). A
# constraint value carrying one of these would otherwise crash the byte-builders.
_ILLEGAL_CELL_CHARS = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")
# Leading chars Excel/Sheets treat as the start of a formula. Prefixing a value
# that begins with one of these neutralizes CSV/formula injection (e.g. a
# constraint Carrier of "=cmd|'/c calc'!A1" must not become a live formula in the
# downloaded report).
_FORMULA_LEAD = ("=", "+", "-", "@")


def _sanitize_cell(val):
    """Make a value safe to write into an xlsx/docx cell.

    Strips illegal control characters and neutralizes leading formula characters
    (prefixing a zero-width-free apostrophe-style guard) so a hostile constraint
    string can neither crash the writer nor execute as a formula when the user
    opens the downloaded report. Non-strings pass through unchanged.
    """
    if not isinstance(val, str):
        return val
    s = _ILLEGAL_CELL_CHARS.sub("", val)
    if s[:1] in _FORMULA_LEAD:
        # A leading apostrophe forces Excel/Sheets to treat the cell as text.
        s = "'" + s
    return s


def _sanitize_frame(df: pd.DataFrame) -> pd.DataFrame:
    """Apply _sanitize_cell to every object/string cell of a DataFrame copy."""
    out = df.copy()
    for col in out.columns:
        if out[col].dtype == object:
            out[col] = out[col].map(_sanitize_cell)
    return out


def build_analysis_workbook_bytes(report: dict,
                                  constraints_after: Optional[list] = None) -> bytes:
    """Multi-sheet, styled .xlsx of a diagnose_constraints report.

    Sheets: Diagnosis Summary, Scope Volume (available vs requested), Issues
    (over-subscription / tiny pools / dead scopes), and — when a corrected set is
    supplied — Corrected Constraints (clean template columns, no helper keys).
    """
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    hdr_fill = PatternFill("solid", fgColor=_REPORT_NAVY)
    hdr_font = Font(bold=True, color="FFFFFF", size=11)
    thin = Side(style="thin", color="D0D0D0")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    wrap = Alignment(wrap_text=True, vertical="top")

    def _frame(rows, columns):
        df = pd.DataFrame(rows, columns=columns) if rows else pd.DataFrame(columns=columns)
        return _sanitize_frame(df)

    # --- Summary sheet ---
    over = report.get("over_subscribed_scopes", []) or []
    tiny = report.get("tiny_pools", []) or []
    dead = report.get("dead_scopes", {}) or {}
    fixable = dead.get("fixable", []) or []
    acceptable = dead.get("acceptable", []) or []
    summary_rows = [
        {"Metric": "Total rules analyzed", "Value": report.get("rule_count", 0)},
        {"Metric": "Over-subscribed scopes", "Value": len(over)},
        {"Metric": "Tiny pools (few containers, many rules)", "Value": len(tiny)},
        {"Metric": "Dead-scope rules (fixable / typos)", "Value": len(fixable)},
        {"Metric": "Dead-scope rules (acceptable / out-of-scope)", "Value": len(acceptable)},
        {"Metric": "Recommended fixes", "Value": len(report.get("recommended_fixes", []) or [])},
    ]

    scope_rows = []
    for s in report.get("scopes", []) or []:
        scope_rows.append({
            "Scope": _scope_label(s.get("port"), s.get("category")),
            "Available": s.get("available_containers"),
            "% Requested": s.get("percent_requested"),
            "Fixed Requested": s.get("fixed_requested"),
            "Total Requested": s.get("total_requested_containers"),
            "Rules": s.get("rule_count"),
            "Over-subscribed": "YES" if s.get("over_subscribed") else "",
        })

    issue_rows = []
    for s in over:
        issue_rows.append({"Issue": "Over-subscription",
                           "Scope": _scope_label(s.get("port"), s.get("category")),
                           "Available": s.get("available_containers"),
                           "Requested": s.get("total_requested_containers"),
                           "Rules": s.get("rule_count"),
                           "Detail": "Caps + percents exceed the pool; lower-priority rules starve."})
    for t in tiny:
        issue_rows.append({"Issue": "Tiny pool",
                           "Scope": _scope_label(t.get("port"), t.get("category")),
                           "Available": t.get("available_containers"),
                           "Requested": "", "Rules": t.get("rule_count"),
                           "Detail": f"Carriers present: {', '.join(t.get('carriers_present') or []) or 'none'}."})
    for d in fixable:
        issue_rows.append({"Issue": "Dead scope (fixable)",
                           "Scope": _scope_label(d.get("Port"), d.get("Category")),
                           "Available": 0, "Requested": "",
                           "Rules": 1,
                           "Detail": f"Dead value(s): {', '.join(d.get('dead_dimensions') or [])}. Carrier {d.get('Carrier','?')}, Lane {d.get('Lane','-')}."})

    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        _frame(summary_rows, ["Metric", "Value"]).to_excel(writer, index=False, sheet_name="Diagnosis Summary")
        _frame(scope_rows, ["Scope", "Available", "% Requested", "Fixed Requested",
                            "Total Requested", "Rules", "Over-subscribed"]).to_excel(
            writer, index=False, sheet_name="Scope Volume")
        _frame(issue_rows, ["Issue", "Scope", "Available", "Requested", "Rules", "Detail"]).to_excel(
            writer, index=False, sheet_name="Issues")
        if constraints_after is not None:
            clean = [c for c in constraints_after if not c.get("_problems")]
            _sanitize_frame(constraints_to_dataframe(clean)).to_excel(
                writer, index=False, sheet_name="Corrected Constraints")

        wb = writer.book
        widths = {
            "Diagnosis Summary": [42, 12],
            "Scope Volume": [22, 12, 13, 16, 16, 8, 16],
            "Issues": [22, 18, 11, 11, 8, 60],
            "Corrected Constraints": [12] * len(CONSTRAINT_COLUMNS),
        }
        for name in wb.sheetnames:
            ws = wb[name]
            for cell in ws[1]:
                cell.fill = hdr_fill
                cell.font = hdr_font
                cell.border = border
                cell.alignment = Alignment(horizontal="center", vertical="center")
            ws.freeze_panes = "A2"
            ws.row_dimensions[1].height = 20
            for i, w in enumerate(widths.get(name, []), start=1):
                ws.column_dimensions[get_column_letter(i)].width = w
            # Color the "Over-subscribed"/"Issue" column for quick scanning.
            hdrs = [c.value for c in ws[1]]
            flag_col = None
            if "Over-subscribed" in hdrs:
                flag_col = hdrs.index("Over-subscribed") + 1
            elif "Issue" in hdrs:
                flag_col = hdrs.index("Issue") + 1
            for r in range(2, ws.max_row + 1):
                for c in range(1, ws.max_column + 1):
                    ws.cell(r, c).border = border
                    ws.cell(r, c).alignment = wrap
                if flag_col:
                    val = str(ws.cell(r, flag_col).value or "")
                    fill = None
                    if val == "YES" or val.startswith("Over"):
                        fill = _REPORT_STATUS_FILL["over"]
                    elif val.startswith("Tiny"):
                        fill = _REPORT_STATUS_FILL["tiny"]
                    elif val.startswith("Dead"):
                        fill = _REPORT_STATUS_FILL["dead"]
                    if fill:
                        ws.cell(r, flag_col).fill = PatternFill("solid", fgColor=fill)
    buf.seek(0)
    return buf.getvalue()


def build_analysis_report_docx_bytes(report: dict,
                                     repair: Optional[dict] = None) -> Optional[bytes]:
    """Narrative Word report of a diagnosis (+ optional repair). None if python-docx absent.

    The guarded import means a clean deploy without python-docx still works — the
    caller produces the Excel workbook and notes that the Word version was skipped.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except Exception:
        return None

    navy = RGBColor(0x1F, 0x4E, 0x78)

    def _shade(cell, hexc):
        tcPr = cell._tc.get_or_add_tcPr()
        sh = OxmlElement("w:shd")
        sh.set(qn("w:val"), "clear")
        sh.set(qn("w:fill"), hexc)
        tcPr.append(sh)

    def _h(doc, text, level=1):
        p = doc.add_heading(text, level=level)
        for r in p.runs:
            r.font.color.rgb = navy
        return p

    def _table(doc, headers, rows, widths=None):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Light Grid Accent 1"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, htext in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(htext)
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9)
            _shade(cell, _REPORT_NAVY)
        for row in rows:
            cells = t.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = ""
                # Strip control chars lxml rejects (no formula guard needed — a
                # Word run is never evaluated as a formula).
                text = "" if val is None else _ILLEGAL_CELL_CHARS.sub("", str(val))
                r = cells[i].paragraphs[0].add_run(text)
                r.font.size = Pt(8.5)
        if widths:
            for i, w in enumerate(widths):
                for row in t.rows:
                    row.cells[i].width = Inches(w)
        return t

    doc = Document()
    title = doc.add_heading("Constraint Analysis Report", level=0)
    for r in title.runs:
        r.font.color.rgb = navy
    doc.add_paragraph(
        "Automated diagnosis of the current constraint working set against the "
        "loaded data. Percentages are computed against the original scope pool; "
        "higher Priority Score is processed first."
    ).runs[0].italic = True

    over = report.get("over_subscribed_scopes", []) or []
    tiny = report.get("tiny_pools", []) or []
    fixable = (report.get("dead_scopes", {}) or {}).get("fixable", []) or []
    acceptable = (report.get("dead_scopes", {}) or {}).get("acceptable", []) or []

    _h(doc, "Summary", 1)
    _table(doc, ["Metric", "Value"], [
        ["Total rules analyzed", report.get("rule_count", 0)],
        ["Over-subscribed scopes", len(over)],
        ["Tiny pools (few containers, many rules)", len(tiny)],
        ["Dead-scope rules (fixable / typos)", len(fixable)],
        ["Dead-scope rules (acceptable / out-of-scope)", len(acceptable)],
    ], widths=[3.5, 1.2])

    if over:
        _h(doc, "Over-subscribed scopes", 1)
        doc.add_paragraph(
            "These scopes' fixed caps + percent rules together request more "
            "containers than the pool holds, so lower-priority rules can't be met."
        )
        _table(doc, ["Scope", "Available", "Requested", "Rules"],
               [[_scope_label(s.get("port"), s.get("category")),
                 s.get("available_containers"), s.get("total_requested_containers"),
                 s.get("rule_count")] for s in over],
               widths=[2.2, 1.2, 1.2, 0.9])

    if tiny:
        _h(doc, "Tiny pools", 1)
        doc.add_paragraph(
            "A scope with very few containers but many rules — the lower-priority "
            "rules can never bind because higher ones consume the handful."
        )
        _table(doc, ["Scope", "Containers", "Rules", "Carriers present"],
               [[_scope_label(t.get("port"), t.get("category")),
                 t.get("available_containers"), t.get("rule_count"),
                 ", ".join(t.get("carriers_present") or []) or "none"] for t in tiny],
               widths=[2.0, 1.1, 0.8, 2.2])

    if fixable:
        _h(doc, "Dead-scope rules (fixable)", 1)
        doc.add_paragraph(
            "These rules target a scope value not present in the data — likely a "
            "typo or a stale code. Correct or drop them."
        )
        _table(doc, ["Carrier", "Port", "Lane", "Dead value(s)"],
               [[d.get("Carrier", "?"), d.get("Port", "-"), d.get("Lane", "-"),
                 ", ".join(d.get("dead_dimensions") or [])] for d in fixable],
               widths=[1.2, 1.0, 1.2, 2.0])

    if repair and repair.get("changes"):
        _h(doc, "Applied corrections (staged for review)", 1)
        s = repair.get("summary", {})
        doc.add_paragraph(
            f"The corrected set rescaled {s.get('rescaled', 0)} percent rule(s), "
            f"dropped {s.get('dropped', 0)} dead rule(s), and collapsed "
            f"{s.get('collapsed', 0)} redundant tiny-pool rule(s) — "
            f"{s.get('kept', 0)} of {s.get('original_count', 0)} rules kept. "
            "These changes are staged for review, not applied."
        )

    _h(doc, "Bottom line", 1)
    doc.add_paragraph(
        "The constraint engine is working correctly; the issues above are in the "
        "constraint list itself. Use the corrected set (staged in the review panel) "
        "to resolve over-subscription and remove rules that can't bind, then Apply "
        "after reviewing. Full per-scope and per-rule data is in the companion "
        "workbook."
    )

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.getvalue()


def constraints_from_dataframe(df, origin: str = "uploaded",
                               valid_carriers: Optional[set] = None) -> list:
    """Convert a processed constraint DataFrame into the working-set list.

    Each row is normalized to the full column schema, validated, and tagged with
    its origin ('uploaded' or 'assistant') so the assistant — and the user — can
    see which rules came from a file vs were drafted in chat.
    """
    if df is None or len(df) == 0:
        return []
    rows = []
    for rec in df.to_dict("records"):
        row = _normalize_constraint(rec)
        row["_problems"] = validate_constraint(row, valid_carriers)
        row["_origin"] = origin
        rows.append(row)
    return rows
