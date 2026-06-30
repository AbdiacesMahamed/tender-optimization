"""
Generate the Tender Optimization Dashboard user guide as a Word document.

Covers three areas:
  1. How to use constraints
  2. How the optimization works
  3. How to use the agent (AI) chatbot

Run:
    .venv/Scripts/python.exe docs/generate_user_guide.py
Output:
    docs/Tender_Optimization_User_Guide.docx
"""
from __future__ import annotations

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

# ----- palette -----
NAVY = RGBColor(0x1F, 0x37, 0x64)
BLUE = RGBColor(0x2E, 0x5C, 0x9E)
GREY = RGBColor(0x55, 0x55, 0x55)
LIGHT = "DCE6F1"
HEADER_FILL = "1F3764"


def shade(cell, hex_fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): hex_fill,
    })
    tc_pr.append(shd)


def set_cell_text(cell, text, *, bold=False, color=None, size=9, white=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    elif color is not None:
        run.font.color.rgb = color


def add_table(doc, headers, rows, *, col_widths=None, font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, white=True, size=font_size)
        shade(hdr[i], HEADER_FILL)
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=font_size)
            if r_idx % 2 == 1:
                shade(cells[i], LIGHT)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def h1(doc, text):
    p = doc.add_heading(level=1)
    run = p.add_run(text)
    run.font.color.rgb = NAVY
    run.font.size = Pt(18)
    return p


def h2(doc, text):
    p = doc.add_heading(level=2)
    run = p.add_run(text)
    run.font.color.rgb = BLUE
    run.font.size = Pt(14)
    return p


def h3(doc, text):
    p = doc.add_heading(level=3)
    run = p.add_run(text)
    run.font.color.rgb = BLUE
    run.font.size = Pt(12)
    return p


def body(doc, text, *, italic=False, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(6)
    return p


def bullet(doc, text, *, level=0, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Inches(0.5 + 0.25 * level)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p


def numbered(doc, text, *, bold_lead=None):
    p = doc.add_paragraph(style="List Number")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
    p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p


def mono(doc, text):
    """Monospace code-ish block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(6)
    for line in text.split("\n"):
        run = p.add_run(line + "\n")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return p


def callout(doc, label, text, fill="FFF2CC"):
    """A single-cell shaded note box."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    shade(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(f"{label}  ")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    doc.add_paragraph()
    return table


# ==================================================================
#  BUILD
# ==================================================================
doc = Document()

# Base style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)

# ---------- Cover ----------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(120)
r = title.add_run("Carrier Tender Optimization Dashboard")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("User Guide")
r.font.size = Pt(20)
r.font.color.rgb = BLUE

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("Constraints  •  Optimization  •  AI Assistant")
r.font.size = Pt(13)
r.font.color.rgb = GREY
r.italic = True

tag = doc.add_paragraph()
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
tag.paragraph_format.space_before = Pt(40)
r = tag.add_run(
    "A Streamlit application that optimizes carrier allocations for inbound "
    "drayage by analyzing container volumes, carrier rates, and performance scores."
)
r.font.size = Pt(11)
r.italic = True
r.font.color.rgb = GREY

doc.add_page_break()

# ---------- TOC-ish intro ----------
h1(doc, "About This Guide")
body(doc,
     "This guide explains the three capabilities you will use most in the Tender "
     "Optimization Dashboard:")
bullet(doc, "lock specific carrier allocations so the optimizer cannot change them.",
       bold_lead="1. Constraints — ")
bullet(doc, "the four allocation scenarios and how the Optimized (LP) strategy decides "
            "who gets which containers.", bold_lead="2. Optimization — ")
bullet(doc, "a Bedrock-powered chatbot in the sidebar that analyzes data, prices "
            "carrier flips, and drafts constraints from plain English.",
       bold_lead="3. AI Assistant — ")
body(doc, "")

h2(doc, "Before You Start: Loading Data")
body(doc, "Run the dashboard with:")
mono(doc, "streamlit run dashboard.py")
body(doc, "Then upload your files from the upload section at the top of the page:")
add_table(
    doc,
    ["File", "Required", "What it provides"],
    [
        ["GVT File", "Yes", "Container movement data — ports, facilities, carriers (SCAC), container numbers, week, category."],
        ["Rate File", "Yes", "Carrier rates by lane (Base Rate per container, plus the lane/port/facility lookup key)."],
        ["Performance File", "Optional", "Carrier performance scores (0–1) by week. Needed for the Performance scenario."],
        ["Constraints File", "Optional", "Operational rules that lock allocations (Excel). Covered in Part 1."],
    ],
    col_widths=[1.4, 1.0, 4.2],
)
callout(doc, "Tip:",
        "Everything downstream — filters, metrics, scenarios, and the AI assistant — "
        "operates on the data you load here. You can also draft constraints with the "
        "assistant before any data is loaded.")

doc.add_page_break()

# ==================================================================
#  PART 1 — CONSTRAINTS
# ==================================================================
h1(doc, "Part 1 — How to Use Constraints")
body(doc,
     "Constraints lock specific allocations so they are not changed by the optimization "
     "scenarios. You provide them as an Excel file (a template lives at "
     "docs/constraint_template.xlsx), or you let the AI assistant draft them for you "
     "(Part 3). Each row in the file is one rule.")

h2(doc, "1.1 The Constraint File Columns")
add_table(
    doc,
    ["Column", "Required", "Description"],
    [
        ["Priority Score", "Yes", "Processing order. Higher = processed first. When two rules compete for the same containers, the higher priority wins."],
        ["Carrier", "Yes*", "The carrier (SCAC) to assign volume to. Required for Maximum, Minimum, and Excluded FC rules. It is the assignment TARGET, not a filter."],
        ["Category", "No", "Scope filter — business category (e.g. FBA FCL, Retail CD). Blank = all."],
        ["Lane", "No", "Scope filter — lane code (e.g. USLAXIUSF). Blank = all."],
        ["Port", "No", "Scope filter — discharged port (e.g. LAX, BAL). Blank = all."],
        ["Week Number", "No", "Scope filter — a specific week. Blank = all weeks."],
        ["Terminal", "No", "Scope filter — port terminal. Blank = all."],
        ["SSL", "No", "Scope filter — steamship line code. Blank = all."],
        ["Vessel", "No", "Scope filter — vessel name. Blank = all."],
        ["Maximum Container Count", "No", "Hard cap on containers for this carrier in scope."],
        ["Minimum Container Count", "No", "Floor — guarantee at least this many containers."],
        ["Percent Allocation", "No", "Percent (0–100) of matching containers to assign to this carrier."],
        ["Excluded FC", "No", "Facility code where this carrier is banned from receiving volume."],
    ],
    col_widths=[1.7, 0.9, 4.0],
)
callout(doc, "*Conditional:",
        "Carrier is required whenever you set a Maximum, Minimum, or Excluded FC rule. "
        "A rule must also do something — set at least one of Maximum / Minimum / "
        "Percent Allocation / Excluded FC, or it has no effect.")

h2(doc, "1.2 The Four Constraint Types")

h3(doc, "Maximum Container Count")
body(doc, "Caps how many containers a carrier can receive in the matched scope.")
bullet(doc, "Containers up to the cap go to the locked Constrained Table.")
bullet(doc, "The carrier is added to the exclusion list so scenarios won't give it more volume.")
bullet(doc, "Excess containers are NOT deleted — they stay in the unconstrained pool for other carriers.")
body(doc, "Example — XPDR gets at most 200 containers total:")
mono(doc, "Priority Score: 100\nCarrier: XPDR\nMaximum Container Count: 200")

h3(doc, "Minimum Container Count")
body(doc, "Guarantees a carrier receives at least this many containers from the matching group.")
body(doc, "Example — EFGH gets at least 30 containers at port LAX:")
mono(doc, "Priority Score: 80\nCarrier: EFGH\nPort: LAX\nMinimum Container Count: 30")

h3(doc, "Percent Allocation")
body(doc, "Assigns a percentage (0–100) of the matching containers to the carrier. If the "
          "pool is consumed by higher-priority rules, percent constraints fall back to the "
          "remaining containers.")
body(doc, "Example — IJKL gets 40% of Retail CD containers at BAL in week 10:")
mono(doc, "Priority Score: 70\nCarrier: IJKL\nCategory: Retail CD\nPort: BAL\nWeek Number: 10\nPercent Allocation: 40")

h3(doc, "Excluded FC (Facility Exclusion)")
body(doc, "Bans a carrier from receiving ANY containers at a specific facility, across all "
          "scenarios. If the carrier already holds containers there, they are reallocated to "
          "capable alternative carriers (checked against the rate data).")
body(doc, "Example — QRST is banned from facility HGR6 entirely:")
mono(doc, "Priority Score: 50\nCarrier: QRST\nExcluded FC: HGR6")
body(doc, "You can combine a cap with an exclusion — e.g. MNOP gets max 100 containers but "
          "never at facility IUSF:")
mono(doc, "Priority Score: 60\nCarrier: MNOP\nMaximum Container Count: 100\nExcluded FC: IUSF")

h2(doc, "1.3 How Scope Filters Work")
body(doc, "Filters narrow which containers a rule applies to. They stack — if you specify "
          "several, ALL must match. Leave a column blank to match every value for that "
          "dimension.")
add_table(
    doc,
    ["Filters specified", "Containers affected"],
    [
        ["None", "All containers for that carrier"],
        ["Category only", "Only containers in that category"],
        ["Category + Lane", "Only containers in that category AND lane"],
        ["Category + Lane + Week", "Only containers matching all three"],
        ["Port only", "Only containers at that port"],
    ],
    col_widths=[2.6, 4.0],
)

h2(doc, "1.4 Processing Order")
numbered(doc, "Constraints are sorted by Priority Score, highest first.")
numbered(doc, "Each constraint is processed in order; matching containers move from the unconstrained to the constrained pool.")
numbered(doc, "Once a container is constrained, a lower-priority rule cannot claim it.")
numbered(doc, "After the file rules, peel pile allocations are applied (see 1.6).")
numbered(doc, "Remaining unconstrained containers are handed to the scenario optimizer.")

h2(doc, "1.5 How Constraints Affect Each Scenario")
add_table(
    doc,
    ["Scenario", "Constrained containers", "Unconstrained containers"],
    [
        ["Current Selection", "Shown locked (unchanged)", "Shown as-is"],
        ["Performance", "Shown locked (unchanged)", "Reallocated to highest performer"],
        ["Cheapest Cost", "Shown locked (unchanged)", "Reallocated to cheapest carrier"],
        ["Optimized (LP)", "Shown locked (unchanged)", "Optimized via LP + historical limits"],
    ],
    col_widths=[1.7, 2.4, 2.5],
)
body(doc, "The total cost in the cost cards always includes both constrained and "
          "unconstrained costs. The constraint summary panel shows which rules were "
          "applied, how many containers each affected, and eligibility / scope / 'why' "
          "diagnostics.")

h2(doc, "1.6 Peel Pile Allocations")
body(doc, "A peel pile is a group of containers from the same Vessel + Week + Discharged "
          "Port + Terminal that is large enough to justify a dedicated carrier assignment. "
          "The qualifying size depends on the port and terminal: the Pacific-Northwest "
          "ports (Tacoma/TIW and Seattle/SEA) use per-terminal limits (Washington United 40, "
          "Husky 45, Terminal 18 and Terminal 5 30, and an 80-container default for any other "
          "PNW terminal), while Oakland (OAK) uses 30. Other ports do not surface peel piles. "
          "The dashboard finds these automatically.")
h3(doc, "Workflow")
numbered(doc, "Scroll to the Peel Pile Analysis section under the analysis table.")
numbered(doc, "Review the qualifying vessel groups (size threshold depends on the port/terminal).")
numbered(doc, "Select a group from the dropdown.")
numbered(doc, "Pick one or more carriers in the multiselect.")
numbered(doc, "Click Add to Queue (fast — queues without recalculating). Repeat for more groups.")
numbered(doc, "Click Apply All to lock all queued assignments and trigger a full recalculation.")
numbered(doc, "Use Clear Queue to discard pending assignments, or Clear All to remove all peel pile assignments.")
body(doc, "When you assign multiple carriers, containers split as evenly as possible — the "
          "first carrier(s) get any extra (e.g. 47 across XPDR + ABCD → 24 / 23). Peel pile "
          "allocations are applied after the file rules, so containers already claimed by a "
          "higher-priority constraint are not available, and peel pile carriers are excluded "
          "from further optimization. Use Download Peel Pile to export a CSV of all groups.")

h2(doc, "1.7 Tips")
bullet(doc, "Use high priority scores (90–100) for hard business rules.")
bullet(doc, "Use lower scores (50–70) for soft preferences.")
bullet(doc, "Combine Maximum + Excluded FC to cap a carrier and ban it from certain facilities.")
bullet(doc, "Use Percent Allocation for proportional splits.")
bullet(doc, "Check the constraint summary to confirm how many containers each rule moved.")

doc.add_page_break()

# ==================================================================
#  PART 2 — OPTIMIZATION
# ==================================================================
h1(doc, "Part 2 — How the Optimization Works")
body(doc, "The dashboard offers four allocation scenarios. You pick one to see how container "
          "assignments — and total cost — would change. Scenarios only run on the "
          "unconstrained pool; constrained and peel pile allocations stay locked.")

h2(doc, "2.1 The Four Scenarios")
add_table(
    doc,
    ["Scenario", "What it does", "Use case"],
    [
        ["Current Selection", "Shows the data exactly as in the GVT file — no optimization.", "Baseline / current state."],
        ["Performance", "Assigns ALL volume in each Lane+Week+Category group to the highest-performing carrier.", "Maximize service quality."],
        ["Cheapest Cost", "Assigns ALL volume in each group to the carrier with the lowest Base Rate.", "Minimize transportation cost."],
        ["Optimized (LP)", "Balances cost and performance with linear programming, capped by historical growth limits.", "Balanced, operationally realistic recommendation."],
    ],
    col_widths=[1.5, 3.4, 1.7],
)

h2(doc, "2.2 The Optimized (LP) Strategy, Step by Step")
body(doc, "This is the flagship scenario. It runs in four steps:")

h3(doc, "Step 1 — Score and rank carriers")
body(doc, "Within each Lane + Week + Category group, every capable carrier gets an "
          "optimization score combining normalized cost and performance:")
mono(doc, "score = (cost_weight x normalized_cost) + (performance_weight x performance_score)")
bullet(doc, "Default weights: 70% cost, 30% performance (configurable; must sum to 100%).")
bullet(doc, "Lower score is better (lower cost is good). Rank 1 = best carrier for that group.")

h3(doc, "Step 2 — Historical volume analysis")
body(doc, "The optimizer calculates each carrier's average market share over the last 5 "
          "weeks of data, so allocations stay grounded in what carriers have actually handled.")

h3(doc, "Step 3 — Apply growth limits")
body(doc, "To prevent unrealistic overnight shifts, a carrier's new allocation is capped at "
          "its historical share plus a maximum growth percentage (default 30%).")
callout(doc, "Example:",
        "A carrier with 25% historical share can grow to at most 25% x 1.3 = 32.5% of the "
        "group's volume.")

h3(doc, "Step 4 — Cascading allocation")
numbered(doc, "Allocate to the Rank 1 carrier up to its growth limit.")
numbered(doc, "Cascade the remaining volume to Rank 2, then Rank 3, and so on.")
numbered(doc, "If volume still remains after every carrier hits its limit, it goes to Rank 1.")

h2(doc, "2.3 Weight Configuration")
add_table(
    doc,
    ["Weight", "Default", "Effect"],
    [
        ["Cost Weight", "70%", "Higher = prioritize lower cost."],
        ["Performance Weight", "30%", "Higher = prioritize better performance."],
    ],
    col_widths=[2.0, 1.2, 3.4],
)

h2(doc, "2.4 Reading the Carrier Flips Column")
body(doc, "Every scenario export includes a Carrier Flips column that traces how containers "
          "moved between carriers, in the form:")
mono(doc, "Had X -> [Changes] -> Now Y")
bullet(doc, "original containers before optimization.", bold_lead="Had X — ")
bullet(doc, "gained N containers from CARRIER.", bold_lead="From CARRIER (+N) — ")
bullet(doc, "lost N containers to CARRIER.", bold_lead="Lost N -> To CARRIER (-N) — ")
bullet(doc, "final container count after the scenario.", bold_lead="Now Y — ")
body(doc, "Examples:")
mono(doc,
     "Had 4 -> From RKNE (+8) + XPDR (+3), Lost 2 -> To FROT (-2) -> Now 15\n"
     "Had 10 (kept all) -> Now 10\n"
     "Had 0 -> From ATMI (+5) -> Now 5")
body(doc, "When container IDs are enabled, the actual container numbers are shown alongside "
          "the counts.")

h2(doc, "2.5 Standard Export Columns")
add_table(
    doc,
    ["Column", "Meaning"],
    [
        ["Carrier Flips", "Container movement trace (see 2.4)."],
        ["Container Numbers", "Comma-separated container IDs assigned to the carrier."],
        ["NEW SCAC", "The carrier assigned in this scenario (renamed from the original Dray SCAC(FL))."],
        ["Discharged Port", "Port code (OAK, LAX, BAL, ...)."],
        ["Category", "Container category type."],
        ["Week Number", "Week of the allocation."],
    ],
    col_widths=[1.8, 4.6],
)
callout(doc, "Note:",
        "NEW SCAC is named to distinguish the optimized assignment from the original "
        "Dray SCAC(FL), so you can compare them directly in the Carrier Flip Analysis.")

h2(doc, "2.6 Metrics Shown")
add_table(
    doc,
    ["Metric", "How it's calculated"],
    [
        ["Total Cost", "Sum of (Base Rate x Container Count)."],
        ["Total Containers", "Sum of all containers."],
        ["Average Rate", "Total Cost / Total Containers."],
        ["Potential Savings", "Current Cost - Optimized Cost."],
        ["Savings %", "Potential Savings / Current Cost x 100."],
    ],
    col_widths=[1.8, 4.6],
)

h2(doc, "2.7 Recommended Workflow")
numbered(doc, "Start with Current Selection to understand the baseline.")
numbered(doc, "Compare Cheapest and Performance to see the cost/quality trade-off.")
numbered(doc, "Use Optimized for a balanced, growth-limited recommendation.")
numbered(doc, "Apply constraints for operational requirements that must hold.")
numbered(doc, "Review the Carrier Flips column to understand the volume movements.")

doc.add_page_break()

# ==================================================================
#  PART 3 — AI ASSISTANT
# ==================================================================
h1(doc, "Part 3 — How to Use the AI Assistant (Chatbot)")
body(doc, "The Tender Assistant is a Bedrock-powered chatbot embedded in the dashboard "
          "sidebar. It is connected to the data you have loaded and can analyze it, price "
          "carrier flips, and draft constraints from plain English. Find it under the "
          "\"🤖 Tender Assistant\" header in the left sidebar.")

h2(doc, "3.1 What It Can Do")
bullet(doc, "the loaded data — carriers, lanes, ports, weeks, cost, and performance.",
       bold_lead="Analyze ")
bullet(doc, "a flip read-only — \"what does flipping these containers to ATMI cost?\" — "
            "with current vs. new cost, the delta and %, and which containers can't be "
            "priced.", bold_lead="Price ")
bullet(doc, "new allocation constraints from a plain-English description.",
       bold_lead="Generate ")
bullet(doc, "or remove the constraints it has staged.", bold_lead="Edit ")
callout(doc, "Important:",
        "The assistant SIMULATES and DRAFTS only. It cannot dispatch, book, or change any "
        "real allocation, and it cannot apply constraints by itself — you review and click "
        "Apply or Download. It will never quote a cost or saving that didn't come from a "
        "tool result.", fill="FCE4D6")

h2(doc, "3.2 The Tools Behind It")
body(doc, "When you ask a question, the assistant chooses among these tools and grounds its "
          "answer in the real data:")
add_table(
    doc,
    ["Tool", "What it answers"],
    [
        ["analyze_data", "Summaries and rankings: overview, by carrier/lane/port/category/week, cheapest, most expensive, performance."],
        ["describe_selection", "What a selection holds — container count, which carriers, lanes, current cost."],
        ["simulate_flip", "Re-prices a selection as if flipped to a target carrier; headline totals + per-lane breakdown."],
        ["compare_carriers", "Prices the same selection under several carriers, ranked cheapest-first."],
        ["lane_rate_options", "For the lanes in a selection, which carriers have a published rate, cheapest first."],
        ["flip_report", "Per-container audit: each container's current carrier, old rate, new rate, and savings."],
        ["preview_constraint_scope", "How many containers a proposed constraint's filters would match."],
        ["generate_constraints", "Validates and stages a batch of proposed constraint rows."],
        ["edit_constraints", "Updates, deletes, or adds rows in the staged constraint list."],
    ],
    col_widths=[1.9, 4.5],
)

h2(doc, "3.3 Example Prompts")
h3(doc, "Analysis")
bullet(doc, "\"Give me an overview of the loaded data.\"")
bullet(doc, "\"Which carriers are cheapest on average?\"")
bullet(doc, "\"Show containers by port.\"")
bullet(doc, "\"Rank carriers by performance.\"")
h3(doc, "Pricing a flip")
bullet(doc, "\"What does flipping all BAL containers to ATMI cost?\"")
bullet(doc, "\"Is Cargomatic or Forrest cheaper for the USLAXIUSF lane?\"")
bullet(doc, "\"Who can serve the HGR6 lane and at what rate?\"")
bullet(doc, "\"Give me a per-container flip report for week 9 to XPDR.\"")
h3(doc, "Drafting constraints")
bullet(doc, "\"Cap XPDR at 200 containers.\"")
bullet(doc, "\"Give ATMI 40% of FBA FCL volume at LAX.\"")
bullet(doc, "\"Ban FRQT from facility BWI4.\"")
bullet(doc, "\"Change that XPDR cap to 150 and bump its priority to 95.\"")
callout(doc, "Carrier names:",
        "You can use either the 4-letter SCAC code or the carrier name (e.g. \"Cargomatic\" "
        "= ATMI). If a name is ambiguous, the assistant will ask rather than guess.")

h2(doc, "3.4 Reviewing, Applying, and Downloading Constraints")
body(doc, "When the assistant proposes constraints, they appear in an editable "
          "\"📋 Proposed constraints\" table in the sidebar. From there:")
numbered(doc, "Review each row. Rows with validation problems are flagged with the reason.",
         bold_lead="Review — ")
numbered(doc, "Edit any cell directly in the table, or add/remove rows.",
         bold_lead="Edit — ")
numbered(doc, "Click ✅ Apply to inject the valid constraints into the live optimization "
              "(no re-upload needed) — the page recalculates with them included.",
         bold_lead="Apply — ")
numbered(doc, "Click 📥 Download to save them as an .xlsx matching the constraint template "
              "for reuse later.", bold_lead="Download — ")
numbered(doc, "Use 🗑️ Clear proposals to discard the draft, or \"Remove applied AI "
              "constraints\" to pull them back out of the optimization.",
         bold_lead="Clear — ")
body(doc, "Applied AI constraints are merged with any uploaded constraint file and "
          "processed together by priority, exactly like file rows.")

h2(doc, "3.5 Setup & Credentials")
body(doc, "The assistant needs AWS Bedrock credentials, configured in .env:")
add_table(
    doc,
    ["Variable", "Purpose"],
    [
        ["AWS_BEDROCK_API_KEY", "Bedrock bearer token (preferred auth)."],
        ["AWS_accessKeyId / AWS_secretAccessKey", "SigV4 fallback, used automatically if the bearer token is rejected."],
        ["BEDROCK_MODEL_ID", "Inference profile, e.g. us.anthropic.claude-opus-4-8."],
        ["BEDROCK_REGION / AWS_REGION", "Region (defaults to us-east-1)."],
    ],
    col_widths=[2.6, 3.8],
)
bullet(doc, "If the bearer token is stale but valid access keys exist, the client retries "
            "with SigV4 automatically — a stale key won't take the assistant down.")
bullet(doc, "If no credentials are found, the sidebar shows a clear warning; you can still "
            "draft and download constraints in the UI.")

h2(doc, "3.6 Good Practices")
bullet(doc, "Ask the assistant to preview a constraint's scope before applying, to confirm how many containers it touches.")
bullet(doc, "Trust only tool-derived numbers — the assistant is instructed never to invent a cost.")
bullet(doc, "Use \"🧹 New chat\" to reset the conversation and clear staged proposals when starting a fresh task.")
bullet(doc, "Remember unpriced containers: if a carrier has no rate on a lane, those containers can't be priced and are reported as unpriced, not free.")

# ---------- Footer line ----------
doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = foot.add_run("Carrier Tender Optimization Dashboard — Internal use only.")
r.italic = True
r.font.size = Pt(9)
r.font.color.rgb = GREY

# ---------- Save ----------
out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                        "Tender_Optimization_User_Guide.docx")
doc.save(out_path)
print(f"Wrote {out_path}")
