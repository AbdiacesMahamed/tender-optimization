"""File-integrity + adversarial stress for the report byte-builders.

Targets ONLY:
  T.build_analysis_workbook_bytes(report, constraints_after=None)
  T.build_analysis_report_docx_bytes(report, repair=None)

Re-opens every produced xlsx with openpyxl AND pandas.read_excel(sheet_name=None),
and every docx with python-docx Document(); injects hostile cell content; pushes
Excel limits. Prints OBSERVED vs EXPECTED for any anomaly.
Run: .venv/Scripts/python.exe scratch_report_integrity.py
"""
import sys, os, io, traceback, warnings
warnings.filterwarnings("ignore")
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import tests.conftest  # streamlit stub  # noqa: F401
import pandas as pd
import components.chatbot.tools as T
from components.constraints.processor import process_constraints_file

import openpyxl
from openpyxl.utils.exceptions import IllegalCharacterError
from docx import Document

comp = pd.read_pickle("/tmp/comp_real.pkl")
PORTS = list(comp["Discharged Port"].dropna().unique())
CATS = list(comp["Category"].dropna().unique())
CARRIERS = list(comp["Dray SCAC(FL)"].dropna().unique())
print(f"comp loaded: {len(comp)} rows, {len(PORTS)} ports, {len(CATS)} cats, {len(CARRIERS)} carriers")

FINDINGS = []
def rec(tag, detail):
    FINDINGS.append((tag, detail))
    print(f"  !! {tag}: {detail}")

EXPECTED_BASE_SHEETS = ["Diagnosis Summary", "Scope Volume", "Issues"]

def verify_xlsx(xb, label, expect_corrected=False):
    """Re-open with openpyxl AND pandas; assert sheets+header present. Returns ok bool."""
    ok = True
    if not isinstance(xb, bytes) or len(xb) < 50:
        rec("xlsx-empty", f"[{label}] builder returned {type(xb)} len={len(xb) if hasattr(xb,'__len__') else 'NA'}")
        return False
    # openpyxl
    try:
        wb = openpyxl.load_workbook(io.BytesIO(xb))
        names = wb.sheetnames
        for s in EXPECTED_BASE_SHEETS:
            if s not in names:
                rec("xlsx-missing-sheet", f"[{label}] expected sheet '{s}' missing; got {names}")
                ok = False
        if expect_corrected and "Corrected Constraints" not in names:
            rec("xlsx-missing-corrected", f"[{label}] 'Corrected Constraints' expected but missing; got {names}")
            ok = False
        for s in names:
            ws = wb[s]
            # header row must exist (row 1 has at least one non-None cell)
            hdr = [c.value for c in ws[1]]
            if all(v is None for v in hdr):
                rec("xlsx-no-header", f"[{label}] sheet '{s}' has empty header row")
                ok = False
    except Exception as e:
        rec("xlsx-openpyxl-FAIL", f"[{label}] {type(e).__name__}: {e}")
        return False
    # pandas read all sheets
    try:
        sheets = pd.read_excel(io.BytesIO(xb), sheet_name=None)
        if not sheets:
            rec("xlsx-pandas-empty", f"[{label}] pandas read 0 sheets")
            ok = False
    except Exception as e:
        rec("xlsx-pandas-FAIL", f"[{label}] {type(e).__name__}: {e}")
        ok = False
    return ok

def verify_docx(db, label):
    if db is None:
        rec("docx-None", f"[{label}] docx builder returned None (python-docx IS installed here)")
        return False
    if not isinstance(db, bytes) or len(db) < 50:
        rec("docx-empty", f"[{label}] builder returned {type(db)} len={len(db) if hasattr(db,'__len__') else 'NA'}")
        return False
    try:
        doc = Document(io.BytesIO(db))
        # touch every table cell to confirm structure is intact
        ncells = 0
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    _ = cell.text
                    ncells += 1
        _ = len(doc.paragraphs)
        return True
    except Exception as e:
        rec("docx-FAIL", f"[{label}] {type(e).__name__}: {e}")
        return False


def build_and_verify(report, label, repair=None, constraints_after=None, expect_corrected=False):
    """Build both artifacts (capturing any crash) and verify integrity."""
    xb = None
    try:
        xb = T.build_analysis_workbook_bytes(report, constraints_after=constraints_after)
    except Exception as e:
        rec("xlsx-builder-RAISED", f"[{label}] {type(e).__name__}: {e}\n{traceback.format_exc()[:400]}")
    if xb is not None:
        verify_xlsx(xb, label, expect_corrected=expect_corrected)
    db = None
    try:
        db = T.build_analysis_report_docx_bytes(report, repair=repair)
    except Exception as e:
        rec("docx-builder-RAISED", f"[{label}] {type(e).__name__}: {e}\n{traceback.format_exc()[:400]}")
    verify_docx(db, label)
    return xb, db


# =====================================================================
# ATTACK 1: many diagnosis shapes
# =====================================================================
print("\n=== ATTACK 1: diagnosis shapes ===")

# 1a empty set
d_empty = T.diagnose_constraints([], comp)
print("empty -> analyzable:", d_empty.get("analyzable"))
build_and_verify(d_empty, "empty-set")

# 1b explicit not-analyzable shape
build_and_verify({"analyzable": False}, "literal-not-analyzable")
build_and_verify({"analyzable": False, "note": "x"}, "not-analyzable-note")
# fully empty dict (degenerate)
build_and_verify({}, "bare-empty-dict")

# 1c diagnosis with 0 issues: one clean, well-scoped rule
clean = [T._normalize_constraint({
    "Priority Score": 10, "Carrier": CARRIERS[0], "Port": PORTS[0],
    "Maximum Container Count": 999999})]
for w in clean:
    w["_origin"] = "uploaded"; w["_problems"] = T.validate_constraint(w, None)
d_zero = T.diagnose_constraints(clean, comp)
print("0-issue -> over:", len(d_zero.get("over_subscribed_scopes", [])),
      "tiny:", len(d_zero.get("tiny_pools", [])),
      "dead:", len(d_zero.get("dead_scopes", {}).get("fixable", [])))
rep_zero = T.repair_constraints(clean, comp, valid_carriers=set(CARRIERS))
build_and_verify(d_zero, "zero-issues", repair=rep_zero,
                 constraints_after=rep_zero["constraints"], expect_corrected=True)

# 1d dozens of over_subscribed / tiny / dead entries
# Build many over-subscribed: stack big percents on each (port,cat) scope.
many = []
pid = 1
for p in PORTS:
    for cat in CATS[:3]:
        for _ in range(4):  # 4 rules each summing >100%
            many.append(T._normalize_constraint({
                "Priority Score": pid, "Carrier": CARRIERS[pid % len(CARRIERS)],
                "Port": p, "Category": cat, "Percent Allocation": 40}))
            pid += 1
# add dozens of dead (typo'd lanes/vessels)
for k in range(40):
    many.append(T._normalize_constraint({
        "Priority Score": 50 + k, "Carrier": CARRIERS[0],
        "Port": PORTS[0], "Lane": f"ZZZ{k}", "Maximum Container Count": 5}))
# add tiny pools: scope a rare (port,cat) with many rules — find a small one
small_scope = None
gp = comp.groupby(["Discharged Port", "Category"])["Container Count"].sum()
small = gp[gp <= 10]
if len(small):
    sp, sc = small.index[0]
    for j in range(5):
        many.append(T._normalize_constraint({
            "Priority Score": 100 + j, "Carrier": CARRIERS[j % len(CARRIERS)],
            "Port": sp, "Category": sc, "Maximum Container Count": 3}))
for w in many:
    w["_origin"] = "uploaded"; w["_problems"] = T.validate_constraint(w, None)
d_many = T.diagnose_constraints(many, comp)
print("dozens -> rules:", d_many["rule_count"],
      "over:", len(d_many["over_subscribed_scopes"]),
      "tiny:", len(d_many["tiny_pools"]),
      "dead-fix:", len(d_many["dead_scopes"]["fixable"]),
      "dead-ok:", len(d_many["dead_scopes"]["acceptable"]))
rep_many = T.repair_constraints(many, comp, valid_carriers=set(CARRIERS))
build_and_verify(d_many, "dozens-of-issues", repair=rep_many,
                 constraints_after=rep_many["constraints"], expect_corrected=True)

# 1e real file
REAL_XLSX = "C:/Users/maabdiac/Downloads/New constraints 5.19 (2).xlsx"
real_ws = None
if os.path.exists(REAL_XLSX):
    try:
        cdf = process_constraints_file(REAL_XLSX)
        real_ws = T.constraints_from_dataframe(cdf, valid_carriers=set(CARRIERS))
        print(f"real file: {len(real_ws)} constraints parsed")
        d_real = T.diagnose_constraints(real_ws, comp)
        print("real -> rules:", d_real["rule_count"],
              "over:", len(d_real["over_subscribed_scopes"]),
              "tiny:", len(d_real["tiny_pools"]),
              "dead-fix:", len(d_real["dead_scopes"]["fixable"]),
              "dead-ok:", len(d_real["dead_scopes"]["acceptable"]))
        rep_real = T.repair_constraints(real_ws, comp, valid_carriers=set(CARRIERS))
        build_and_verify(d_real, "REAL-file", repair=rep_real,
                         constraints_after=rep_real["constraints"], expect_corrected=True)
    except Exception as e:
        rec("real-file-pipeline-RAISED", f"{type(e).__name__}: {e}\n{traceback.format_exc()[:400]}")
else:
    print("REAL file not found, skipping")


# =====================================================================
# ATTACK 3: adversarial cell content (formula injection, huge, control chars)
# =====================================================================
print("\n=== ATTACK 3: adversarial content ===")
INJECT = "=cmd|'/c calc'!A1"
HOSTILE = [
    INJECT, "@SUM(1+1)*cmd", "+1+1", "-1-1", "=1+1",
    "A" * 40000,
    "line1\nline2\ttab", "\r\n\r\n",
    "unicode ‮ rtl \U0001F4A9 emoji 中文",
    "\x00\x01\x07\x1f control",  # control chars incl NUL + bell
    "\x0b\x0c vt ff",
]

def make_report_with_value(v, kind="dead"):
    """Build a diagnose-shaped report whose cells carry hostile value v."""
    if kind == "dead":
        return {
            "analyzable": True, "have_data": True, "rule_count": 1,
            "scopes": [{"port": v, "category": v, "available_containers": v,
                        "percent_requested": v, "fixed_requested": v,
                        "total_requested_containers": v, "rule_count": 1,
                        "over_subscribed": True, "rules": [{"Carrier": v}]}],
            "over_subscribed_scopes": [{"port": v, "category": v,
                        "available_containers": v, "total_requested_containers": v,
                        "rule_count": 1}],
            "tiny_pools": [{"port": v, "category": v, "available_containers": v,
                            "rule_count": 3, "carriers_present": [v], "rules": [{"Carrier": v}]}],
            "dead_scopes": {"fixable": [{"Carrier": v, "Port": v, "Lane": v,
                                         "Category": v, "dead_dimensions": [v]}],
                            "acceptable": [{"Carrier": v, "dead_dimensions": [v]}]},
            "recommended_fixes": [{"type": "x", "detail": v}],
        }

for i, v in enumerate(HOSTILE):
    safe_label = repr(v)[:40]
    rep_for_injection = {"summary": {"rescaled": 1, "dropped": 0, "collapsed": 0,
                                     "kept": 1, "original_count": 1}, "changes": [{"index": 0}]}
    # constraints_after carries hostile strings into the Corrected Constraints sheet too
    ca = [T._normalize_constraint({"Priority Score": 1, "Carrier": v, "Lane": v, "Port": v})]
    for w in ca:
        w["_problems"] = []  # force it into the clean sheet
    rpt = make_report_with_value(v)
    xb, db = build_and_verify(rpt, f"hostile[{i}]={safe_label}", repair=rep_for_injection,
                              constraints_after=ca, expect_corrected=True)
    # Formula-injection check: is the raw "=..."/"@"/"+"/"-" written unescaped into a cell?
    if xb and v in (INJECT, "@SUM(1+1)*cmd", "+1+1", "-1-1", "=1+1"):
        try:
            wb = openpyxl.load_workbook(io.BytesIO(xb))
            raw_hits = []
            for sn in wb.sheetnames:
                ws = wb[sn]
                for row in ws.iter_rows():
                    for c in row:
                        if isinstance(c.value, str) and c.value == v and v[0] in "=@+-":
                            raw_hits.append(f"{sn}!{c.coordinate}")
            if raw_hits:
                rec("FORMULA-INJECTION-RAW", f"value {v!r} written verbatim (not prefixed/escaped) at {raw_hits[:5]}")
        except Exception as e:
            rec("injection-check-FAIL", f"{v!r}: {e}")


# =====================================================================
# ATTACK 4: Excel limits — >32767-char cell, huge table
# =====================================================================
print("\n=== ATTACK 4: Excel limits ===")
BIG = "B" * 40000  # exceeds 32767 char/cell limit
rpt_big = make_report_with_value(BIG)
# direct openpyxl behavior probe
try:
    wb = openpyxl.Workbook(); ws = wb.active
    ws["A1"] = BIG
    b = io.BytesIO(); wb.save(b)
    # re-open and measure
    wb2 = openpyxl.load_workbook(io.BytesIO(b.getvalue()))
    got = wb2.active["A1"].value
    print(f"openpyxl direct 40k cell: saved & reloaded len={len(got) if got else 0} (limit 32767) -> {'TRUNCATED' if got and len(got)<40000 else 'KEPT FULL'}")
except Exception as e:
    print(f"openpyxl direct 40k cell RAISED: {type(e).__name__}: {e}")

xb_big, db_big = build_and_verify(rpt_big, "40k-char-cells")

# huge table: many columns / rows scope sheet
huge_scopes = [{"port": f"P{i}", "category": f"C{i}", "available_containers": i,
                "percent_requested": i, "fixed_requested": i,
                "total_requested_containers": i, "rule_count": i,
                "over_subscribed": True} for i in range(20000)]
rpt_huge = {"analyzable": True, "rule_count": 20000, "scopes": huge_scopes,
            "over_subscribed_scopes": huge_scopes[:5000], "tiny_pools": [],
            "dead_scopes": {"fixable": [], "acceptable": []}, "recommended_fixes": []}
build_and_verify(rpt_huge, "20k-row-table")


# =====================================================================
# ATTACK 5: docx from huge/edge diagnosis still opens
# =====================================================================
print("\n=== ATTACK 5: docx huge/edge ===")
verify_docx(T.build_analysis_report_docx_bytes(rpt_huge), "docx-huge")
verify_docx(T.build_analysis_report_docx_bytes(make_report_with_value(BIG)), "docx-40k")
verify_docx(T.build_analysis_report_docx_bytes(make_report_with_value(INJECT)), "docx-injection")
verify_docx(T.build_analysis_report_docx_bytes({"analyzable": False}), "docx-not-analyzable")
verify_docx(T.build_analysis_report_docx_bytes({}), "docx-bare-empty")


print("\n" + "=" * 80)
if FINDINGS:
    print(f"TOTAL ANOMALIES: {len(FINDINGS)}")
    from collections import Counter
    for tag, n in Counter(t for t, _ in FINDINGS).most_common():
        print(f"  {tag}: {n}")
else:
    print("NO ANOMALIES — every artifact built and re-opened cleanly.")
