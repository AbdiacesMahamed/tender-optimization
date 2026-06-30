"""Generate a Word document summarizing the PNW (Pacific Northwest) allocation rules.

Source of truth: docs/PNW_RULES.md and components/constraints/pnw_vessel_rules.py.
Where the two disagree, the code wins (Rule 1 is a Max-130 ceiling, no Min floor).
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
GREY = RGBColor(0x55, 0x55, 0x55)


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
    return h


def add_body(doc, text, *, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(size)
    return p


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(it, tuple):
            lead, rest = it
            r = p.add_run(lead)
            r.bold = True
            p.add_run(rest)
        else:
            p.add_run(it)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    return table


def main():
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---- Title ----
    title = doc.add_heading("PNW Allocation Rules", level=0)
    for run in title.runs:
        run.font.color.rgb = NAVY
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = sub.add_run("Pacific Northwest waterfront — Seattle (SEA) + Tacoma (TIW)")
    r.italic = True
    r.font.size = Pt(12)
    r.font.color.rgb = GREY

    add_body(doc,
        "These are standing, always-on rules for the Pacific Northwest waterfront. They are "
        "enforced on every optimization run, on top of (and ahead of) anything a user uploads "
        "or asks the assistant for. Rules 2–5 apply to both PNW ports; Rule 1 is Tacoma-only.")

    # ---- Ports ----
    add_heading(doc, "Ports", 1)
    add_table(doc, ["Code", "Port", "Notes"], [
        ["SEA", "Seattle", "Discharged Port (not a terminal)"],
        ["TIW", "Tacoma", "Discharged Port (not a terminal)"],
    ])
    add_body(doc,
        "Both SEA and TIW are Discharged Ports in the GVT data — not terminals. Most PNW rules "
        "are Port-scoped (Terminal column left blank); the exception is Rule 5 (peel-pile "
        "thresholds), which is per-terminal.")

    # ---- Terminals ----
    add_heading(doc, "PNW Terminals", 2)
    add_table(doc, ["GVT Terminal string", "Terminal", "Port"], [
        ["TRM-T004", "Husky Terminal (Tacoma Terminal 4)", "TIW"],
        ["TRM-TWUT", "Washington United Terminal (WUT)", "TIW"],
        ["TRM-TPCT", "Pierce County Terminal (PCT)", "TIW"],
        ["SSA-T18", "Terminal 18", "SEA"],
        ["TERMINAL 5", "Terminal 5", "SEA"],
    ])

    # ---- Rule summary ----
    add_heading(doc, "Rule Summary", 1)
    add_table(doc, ["#", "Rule", "Type", "Status"], [
        ["0", "Carrier-to-port lockouts (AOYV/RDXY = Seattle, RKNE/HJBT = Tacoma)",
         "Static per-carrier per-port", "Live (config)"],
        ["1", "JB Hunt (HJBT) capped at up to 130 containers/week at Tacoma (TIW), allocated first",
         "Per-carrier, per-week ceiling", "Live (generated rows)"],
        ["2", "No SCAC may take more than 60 containers per vessel",
         "All-carrier, per-vessel cap", "Live (generated rows + safety net)"],
        ["3", "A SCAC may take volume from only 1 vessel at a time",
         "Assignment constraint", "Live (post-alloc pass)"],
        ["4", "If 2+ vessels arrive the same day, a SCAC may take from only 1 of them",
         "Same-day refinement of #3", "Live (post-alloc pass)"],
        ["5", "Peel-pile qualifying thresholds per-terminal (WUT 40, Husky 45, T18/T5 30, default 80)",
         "Per-terminal group-size floor", "Live (config)"],
        ["6", "Empties awareness + prior-week carrier continuity",
         "Planning heuristic (look-back)", "Pending (design)"],
    ])

    # ---- Rule 0 ----
    add_heading(doc, "Rule 0 — Carrier-to-Port Lockouts", 1)
    add_body(doc,
        "A carrier that runs at only one PNW port is locked out (Max 0) of the other port. A "
        "Max 0 lockout allocates nothing to that carrier in scope and prevents the optimizer "
        "from sending it any volume there.")
    add_table(doc, ["Carrier", "SCAC", "Operates at", "Encoded as"], [
        ["Waterfront Logistics", "AOYV", "Seattle only", "Max 0 lockout at TIW (Tacoma)"],
        ["RoadEx", "RDXY", "Seattle only", "Max 0 lockout at TIW (Tacoma)"],
        ["RoadOne Intermodal", "RKNE", "Tacoma only", "Max 0 lockout at SEA (Seattle)"],
        ["JB Hunt", "HJBT", "Tacoma only", "Max 0 lockout at SEA (Seattle)"],
    ])
    add_body(doc,
        "Carriers not listed above (e.g. Forrest / FRQT) are unrestricted in the PNW and may "
        "allocate at both ports. Encoded as prebuilt per-port constraints in "
        "config/port_constraints/SEA.csv and TIW.csv (all Priority Score 100, Port-scoped, Max 0).")

    # ---- Rule 1 ----
    add_heading(doc, "Rule 1 — JB Hunt: up to 130 containers/week at Tacoma", 1)
    add_body(doc,
        "HJBT is capped at up to 130 containers per week at Tacoma (TIW) and is allocated FIRST "
        "(always-on prebuilt priority), so it draws as close to 130 as the week's available TIW "
        "volume allows.")
    add_bullets(doc, [
        ("Scope: ", "Carrier HJBT, Port TIW, per Week Number."),
        ("Amount: ", "Maximum 130 per week. There is NO Min floor."),
        ("Shortfall: ", "A week with fewer than 130 TIW containers simply gives HJBT whatever "
         "is present — no shortfall is forced or flagged."),
    ])
    add_body(doc,
        "Note: an earlier version emitted Min-130 + Max-130 (\"exactly 130\"). The current code "
        "(build_hunt_weekly_rows) emits a Max-130 ceiling only, allocated first.", italic=True)

    # ---- Rule 2 ----
    add_heading(doc, "Rule 2 — Per-Vessel Cap: no SCAC over 60 containers/vessel", 1)
    add_body(doc,
        "No carrier (SCAC) may be allocated more than 60 containers from any single vessel at "
        "PNW ports. This is a cap on every carrier, per vessel — not a rule targeting one named "
        "carrier.")
    add_bullets(doc, [
        ("Scope: ", "PNW ports (SEA + TIW), per Vessel, per SCAC."),
        ("Amount: ", "≤ 60 containers, applied independently to each (vessel, carrier)."),
        ("Enforcement: ", "build_per_vessel_cap_rows() materializes one Max-60 row per (PNW "
         "port, vessel, carrier) present in the data; each becomes a scoped ceiling binding "
         "across both the constrained and unconstrained tables."),
        ("Safety net: ", "enforce_per_vessel_cap_across() runs post-allocation across both "
         "tables, trimming any carrier the optimizer later moved onto a vessel back to ≤ 60. "
         "Excess containers have their carrier cleared so the optimizer re-homes them; volume "
         "is conserved."),
    ])
    add_body(doc,
        "Locked-out carriers (Rule 0) are skipped — a Max-60 \"permission\" row would contradict "
        "their Max-0 lockout.", italic=True)

    # ---- Rules 3 & 4 ----
    add_heading(doc, "Rules 3 & 4 — One Vessel per SCAC", 1)
    add_body(doc,
        "Rule 3: a carrier may take volume from only one vessel at a time — each SCAC's PNW "
        "volume must come from a single vessel, not spread across several.")
    add_body(doc,
        "Rule 4 (refinement of Rule 3): when two or more vessels arrive on the same day (same "
        "Ocean ETA / arrival date at a PNW port), a carrier may take volume from only one of "
        "those same-day vessels.")
    add_bullets(doc, [
        ("Enforcement: ", "enforce_one_vessel_per_carrier_across() runs post-allocation. Within "
         "each (PNW port, arrival-day) group that has 2+ vessels, every carrier is collapsed onto "
         "the single vessel where it already holds the most volume."),
        ("Tie-break: ", "the carrier keeps the vessel with the most containers; ties broken by "
         "vessel name for determinism."),
        ("Displaced volume: ", "containers on the losing same-day vessels have the carrier "
         "cleared so the optimizer re-homes them to an eligible carrier. Volume is conserved "
         "(no rows dropped)."),
    ])
    add_body(doc,
        "These are combinatorial assignment rules no single constraint row can express, which is "
        "why they run as a deterministic post-allocation pass. check_one_vessel_per_carrier() is "
        "the read-only validator used by tests and diagnostics.", italic=True)

    # ---- Rule 5 ----
    add_heading(doc, "Rule 5 — Per-Terminal Peel-Pile Qualifying Thresholds", 1)
    add_body(doc,
        "A peel pile is a Vessel + Week + Discharged Port + Terminal group large enough to "
        "warrant a dedicated carrier assignment. The qualifying threshold (the minimum container "
        "count a group must reach to surface) is per-terminal at PNW, rather than the old global 30.")
    add_table(doc, ["Terminal", "GVT Terminal string", "Threshold"], [
        ["Washington United (WUT)", "TRM-TWUT", "40"],
        ["Husky (Tacoma Terminal 4)", "TRM-T004", "45"],
        ["Terminal 18", "SSA-T18", "30"],
        ["Terminal 5", "TERMINAL 5", "30"],
        ["Pierce County (PCT) / other / blank", "TRM-TPCT, …", "80 (PNW default)"],
    ])
    add_bullets(doc, [
        ("Outside PNW: ", "only Oakland (OAK) keeps a peel-pile threshold (30); every other port "
         "is effectively disabled, so peel piles surface only at PNW and OAK."),
        ("Note: ", "this is a qualifying threshold (which groups show up), not an allocation cap "
         "— once a group qualifies it is split/assigned exactly as before."),
    ])
    add_body(doc,
        "Rule 5 lives in config/peel_pile_thresholds.py (single source of truth), not the "
        "per-port constraint CSVs.", italic=True)

    # ---- Rule 6 ----
    add_heading(doc, "Rule 6 — Empties Awareness + Prior-Week Carrier Continuity (Pending)", 1)
    add_body(doc,
        "Status: PENDING — design captured, not yet implemented. This is a planning heuristic "
        "(a look-back that informs future suggestions), not a hard constraint.", italic=True)
    add_body(doc,
        "Previous/current weeks have already happened — those allocations are locked. They are "
        "still useful as signal: when similar volume recurs, the assistant should suggest the "
        "same SCAC initially allocated the comparable volume, so a carrier that sent 10 "
        "containers into an FC can be planned to pick up 10 empties from there. The goal is "
        "carrier continuity — drop-offs and empty pickups handled by the same carrier — which "
        "reduces stranded empties and repositioning.")
    add_body(doc, "Three inputs drive this:")
    add_bullets(doc, [
        ("Prior-week allocation (look-back): ", "from week W-1 actuals, find which SCAC was "
         "initially allocated each comparable volume slice; that SCAC becomes the preferred carrier."),
        ("Per-terminal volume baseline (continuity target): ", "trailing-3-week average container "
         "count each SCAC handled at each terminal; bias upcoming allocation to land near it. A "
         "soft preference (~15% of a cost-60/perf-25/continuity-15 blend, tunable via "
         "opt_continuity_weight, 0 disables) — never overrides a materially cheaper/better option."),
        ("Empties at the FC: ", "the empties-pending-return report (port-level pivot; USPNW = "
         "TIW + SEA) tells us how many empties are pending per port, bucketed by dwell. High "
         "counts are the pickup opportunity to pair with outbound volume."),
    ])

    # ---- Implementation / wiring ----
    add_heading(doc, "Implementation & Wiring", 1)
    add_table(doc, ["Rule", "Mechanism", "Function"], [
        ["1", "Per-week Max-130 HJBT/TIW row, allocated first", "build_hunt_weekly_rows()"],
        ["2", "Max-60 row per (PNW port, vessel, carrier) + post-alloc safety net",
         "build_per_vessel_cap_rows() + enforce_per_vessel_cap_across()"],
        ["3+4", "Post-alloc pass: collapse each carrier to one same-day vessel; clear the rest",
         "enforce_one_vessel_per_carrier_across()"],
    ])
    add_body(doc,
        "Rules 1–4 live in components/constraints/pnw_vessel_rules.py (pure functions over "
        "DataFrames). Generated rows (Rules 1 + 2) are merged into the always-on front block via "
        "merge_prebuilt_first(constraints_df, data) so they process ahead of user rules. After "
        "apply_constraints_to_data(), the two post-allocation passes run across the combined "
        "constrained + unconstrained tables (the rules bind on a carrier's total PNW volume, not "
        "within either table alone): the Rule 2 safety net first, then Rules 3 & 4.")

    # ---- How to change ----
    add_heading(doc, "How to Change a Rule", 1)
    add_bullets(doc, [
        ("Carrier-port restriction (Rule 0): ", "edit the relevant config/port_constraints/<PORT>.csv "
         "— add/remove a Port-scoped Max 0 row. No engine code changes needed."),
        ("Peel-pile threshold (Rule 5): ", "edit config/peel_pile_thresholds.py "
         "(PNW_TERMINAL_THRESHOLDS, PNW_DEFAULT_THRESHOLD, PORT_THRESHOLDS)."),
        ("Vessel rules (Rules 1–4): ", "edit the config constants at the top of "
         "pnw_vessel_rules.py — HUNT_WEEKLY_MAX (130), PER_VESSEL_MAX (60), PNW_PORTS, "
         "HUNT_SCAC / HUNT_PORT."),
        ("Disable all PNW rules: ", "set \"SEA\" / \"TIW\" to False in the ENABLED dict in "
         "components/constraints/prebuilt.py."),
    ])

    out = Path(__file__).resolve().parents[1] / "docs" / "PNW_Rules_Reference.docx"
    doc.save(str(out))
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
